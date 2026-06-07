"""
make_supplemental_tables.py
Generates iScience_HFpEF_Supplemental_Tables.docx with Table S1 (Key Resources Table)
and Table S2 (GSEA Gene Set Definitions) in iScience / Cell Press format.
"""

import os
import sys

# ── Paths (relative to repo root) ────────────────────────────────────────────
import os as _os
BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))
_os.makedirs(_os.path.join(BASE_DIR, 'figures_out'), exist_ok=True)

PUB  = _os.path.join(BASE_DIR, 'figures_out')
DATA = _os.path.join(BASE_DIR, 'data')
DEG  = _os.path.join(BASE_DIR, 'deg')


try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    sys.exit(f"ERROR: python-docx not available — {e}\nInstall with: pip install python-docx")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
CELL_PRESS_BLUE = RGBColor(0, 70, 127)
LIGHT_BLUE      = RGBColor(173, 198, 224)
WHITE           = RGBColor(255, 255, 255)
BLACK           = RGBColor(0, 0, 0)

FONT_NAME = "Arial"
OUTPUT_PATH = r_os.path.join(BASE_DIR, "iScience_HFpEF_Supplemental_Tables.docx")

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_rgb: RGBColor):
    """Apply solid background shading to a table cell using OxmlElement."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(fill_rgb[0], fill_rgb[1], fill_rgb[2])
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """
    Set borders on a cell.  kwargs keys: top, bottom, left, right
    Each value is a dict with optional keys: sz, val, color, space.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{ edge }")
            params = kwargs[edge]
            tag.set(qn("w:val"),   params.get("val",   "single"))
            tag.set(qn("w:sz"),    str(params.get("sz", 4)))
            tag.set(qn("w:space"), str(params.get("space", 0)))
            tag.set(qn("w:color"), params.get("color", "000000"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run(paragraph, text, bold=False, italic=False,
            font_size=9, color=BLACK, font_name=FONT_NAME):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name      = font_name
    run.font.size      = Pt(font_size)
    run.font.color.rgb = color
    return run


def set_para_spacing(paragraph, space_before=0, space_after=0, line_rule=None):
    pPr  = paragraph._p.get_or_add_pPr()
    spng = OxmlElement("w:spacing")
    spng.set(qn("w:before"), str(int(space_before)))
    spng.set(qn("w:after"),  str(int(space_after)))
    if line_rule:
        spng.set(qn("w:lineRule"), line_rule)
    pPr.append(spng)


def set_column_widths(table, widths_inches):
    """Set column widths for every column in the table."""
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement("w:tcW")
                twips = int(widths_inches[idx] * 1440)  # 1 inch = 1440 twips
                tcW.set(qn("w:w"),    str(twips))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def merge_row_cells(table, row_idx):
    """Merge all cells in a given row (for section headers)."""
    row   = table.rows[row_idx]
    cells = row.cells
    if len(cells) > 1:
        cells[0].merge(cells[-1])
    return cells[0]


def set_table_style_no_border(table):
    """Remove default table border/style and set clean borders."""
    tbl = table._tbl
    # tblPr may already exist; find or create it
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "AAAAAA")
        tblBorders.append(tag)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def create_document():
    doc = Document()

    # Margins: 1 inch on all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Default style font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(9)

    return doc


# ---------------------------------------------------------------------------
# Table S1 — Key Resources Table
# ---------------------------------------------------------------------------

def add_table_s1(doc):
    # --- Table header paragraph ---
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=0, space_after=60)
    add_run(p, "Table S1. Key Resources Table. Related to STAR",
            bold=True, font_size=11, color=CELL_PRESS_BLUE)
    add_run(p, "\u2605",   # ★
            bold=True, font_size=11, color=CELL_PRESS_BLUE)
    add_run(p, " Methods.",
            bold=True, font_size=11, color=CELL_PRESS_BLUE)

    # --- Build table ---
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_style_no_border(table)

    col_widths = [2.2, 2.5, 2.3]  # inches; total ~7 in (within 1-inch margins on letter)
    set_column_widths(table, col_widths)

    # --- Column header row ---
    hdr_row  = table.rows[0]
    hdr_data = ["REAGENT or RESOURCE", "SOURCE", "IDENTIFIER"]
    for idx, cell in enumerate(hdr_row.cells):
        set_cell_shading(cell, CELL_PRESS_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_cell, space_before=40, space_after=40)
        add_run(p_cell, hdr_data[idx], bold=True, font_size=9, color=WHITE)

    # --- Section + data helper ---
    def add_section_header(label):
        row  = table.add_row()
        cell = merge_row_cells(table, len(table.rows) - 1)
        set_cell_shading(cell, LIGHT_BLUE)
        p_sec = cell.paragraphs[0]
        set_para_spacing(p_sec, space_before=40, space_after=40)
        add_run(p_sec, label, bold=True, font_size=9, color=BLACK)

    def add_data_row(col0, col1, col2):
        row   = table.add_row()
        cells = row.cells
        values = [col0, col1, col2]
        for idx, cell in enumerate(cells):
            p_cell = cell.paragraphs[0]
            set_para_spacing(p_cell, space_before=30, space_after=30)
            add_run(p_cell, values[idx], bold=False, font_size=9, color=BLACK)

    # --- Deposited Data ---
    add_section_header("Deposited Data")
    add_data_row(
        "Mouse cardiomyocyte-enriched RNA-seq (WT HFD vs. chow; GSE209548)",
        "NCBI GEO", "GSE209548")
    add_data_row(
        "Mouse whole-heart RNA-seq (HFD+L-NAME model; GSE194151)",
        "NCBI GEO", "GSE194151")
    add_data_row(
        "Human plasma miRNA array (HFpEF vs. healthy; GSE53437)",
        "NCBI GEO", "GSE53437")
    add_data_row(
        "Mouse whole-heart RNA-seq (HFD+mTAC model, L-NAME-free; GSE249409; "
        "cross-model validation dataset)",
        "NCBI GEO", "GSE249409")
    add_data_row(
        "PPARα ligand-binding domain crystal structure",
        "RCSB PDB", "PDB: 1K7L")

    # --- Software and Algorithms ---
    add_section_header("Software and Algorithms")
    sw_rows = [
        ("Python",        "Python Software Foundation",                  "v3.11; https://www.python.org"),
        ("R",             "R Core Team",                                  "v4.3.1; https://www.r-project.org"),
        ("PyDESeq2",      "Muzellec et al., Bioinformatics 2023",         "v0.4; https://github.com/owkin/PyDESeq2"),
        ("gseapy",        "Fang et al.",                                  "v1.0; https://github.com/zqfang/gseapy"),
        ("fgsea",         "Korotkevich et al., 2016, bioRxiv",            "v1.26; https://bioconductor.org/packages/fgsea"),
        ("WGCNA",         "Langfelder & Horvath, BMC Bioinform. 2008",    "v1.72; https://cran.r-project.org/package=WGCNA"),
        ("AutoDock Vina", "Eberhardt et al.",                             "v1.1.2; https://vina.scripps.edu"),
        ("AutoDockTools", "Morris et al.",                                "v1.5.6; https://mgltools.scripps.edu"),
        ("RDKit",         "RDKit Community",                              "2023.03; https://www.rdkit.org"),
        ("NetworkX",      "Hagberg et al.",                               "v3.1; https://networkx.org"),
        ("matplotlib",    "Hunter, J. Comput. Sci. 2007",                 "v3.7; https://matplotlib.org"),
        ("pandas",        "McKinney, 2010",                               "v2.0; https://pandas.pydata.org"),
        ("numpy",         "Harris et al., Nature 2020",                   "v1.24; https://numpy.org"),
        ("scipy",         "Virtanen et al., Nat. Methods 2020",           "v1.11; https://scipy.org"),
        ("GEOparse",      "Guma\u0142a et al.",
         "v2.0; https://github.com/guma44/GEOparse "
         "[used for GEO SOFT file metadata parsing]"),
        ("sklearn (scikit-learn)", "Pedregosa et al., JMLR 2011",
         "v1.3; https://scikit-learn.org [used for PCA in QC figures]"),
        ("Analysis code", "This paper",
         "Available from the lead contact upon request "
         "(zhangliangqing@gdmu.edu.cn)"),
    ]
    for row_data in sw_rows:
        add_data_row(*row_data)

    # --- Other ---
    add_section_header("Other")
    other_rows = [
        ("STRING protein interaction database",
         "Szklarczyk et al., Nucleic Acids Res. 2023",
         "v12.0; https://string-db.org"),
        ("miRDB miRNA target database",
         "Chen & Wang",
         "v5.0; https://mirdb.org"),
        ("TargetScan",
         "McGeary et al., Science 2019",
         "v8.0; https://targetscan.org"),
        ("miRTarBase",
         "Huang et al.",
         "v9.0; https://mirtarbase.cuhk.edu.cn"),
        ("CMAP/L1000 connectivity map",
         "Subramanian et al., Cell 2017",
         "https://clue.io/cmap"),
    ]
    for row_data in other_rows:
        add_data_row(*row_data)

    set_column_widths(table, col_widths)


# ---------------------------------------------------------------------------
# Table S2 — GSEA Gene Set Definitions
# ---------------------------------------------------------------------------

def add_table_s2(doc):
    # --- Page break ---
    p_break = doc.add_paragraph()
    p_break._p.clear()
    run_elem = OxmlElement("w:r")
    br_tag   = OxmlElement("w:br")
    br_tag.set(qn("w:type"), "page")
    run_elem.append(br_tag)
    p_break._p.append(run_elem)

    # --- Table S2 header paragraph ---
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=0, space_after=60)
    add_run(p, "Table S2. GSEA Gene Set Definitions. Related to STAR",
            bold=True, font_size=11, color=CELL_PRESS_BLUE)
    add_run(p, "\u2605",
            bold=True, font_size=11, color=CELL_PRESS_BLUE)
    add_run(p, " Methods.",
            bold=True, font_size=11, color=CELL_PRESS_BLUE)

    # --- Description paragraph ---
    p_desc = doc.add_paragraph()
    set_para_spacing(p_desc, space_before=0, space_after=80)
    add_run(p_desc,
            "Seven a priori gene sets used for pre-ranked GSEA across both mouse HFpEF datasets "
            "(GSE209548 and GSE194151). Gene sets 1\u20133 derived from KEGG pathway gene lists; "
            "gene set 4 from Gene Ontology biological process terms; gene sets 5\u20137 are "
            "investigator-curated from published literature.",
            bold=False, font_size=9, color=BLACK)

    # --- Build table ---
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_style_no_border(table)

    col_widths_s2 = [1.4, 2.3, 0.4, 2.9]  # inches; total ~7 in
    set_column_widths(table, col_widths_s2)

    # --- Column header row ---
    hdr_row  = table.rows[0]
    hdr_data = ["Gene Set", "Source/Annotation", "n", "Gene Members"]
    for idx, cell in enumerate(hdr_row.cells):
        set_cell_shading(cell, CELL_PRESS_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_cell, space_before=40, space_after=40)
        add_run(p_cell, hdr_data[idx], bold=True, font_size=9, color=WHITE)

    # --- Data rows ---
    data_rows = [
        ("FAO_KEGG",
         "KEGG mmu01212 (Fatty acid oxidation)",
         "15",
         "Cpt1b, Acadl, Hadhb, Hadha, Fabp3, Acadvl, Acadm, Cpt2, Hadh, "
         "Slc25a20, Acads, Acsl1, Acsl4, Acsl3, Cpt1a"),
        ("Glycolysis_KEGG",
         "KEGG mmu00010 (Glycolysis/Gluconeogenesis)",
         "9",
         "Hk1, Aldoa, Pfkm, Pkm, Ldhb, Hk2, Gapdh, Ldha, Eno1"),
        ("OxPhos_KEGG",
         "KEGG mmu00190 (Oxidative phosphorylation)",
         "8",
         "Ndufa1, Ndufb1, Sdha, Sdhb, Uqcrc1, Cox4i1, Atp5a1, Cycs"),
        ("Ketone_GO",
         "GO:0006635/GO:0001676 (Fatty acid beta-oxidation/ketone body metabolism)",
         "8",
         "Oxct1, Acat1, Bdh1, Hmgcl, Hmgcs1, Bdh2, Hmgcs2, Oxct2"),
        ("Insulin_Signaling",
         "Investigator-curated (KEGG mmu04910)",
         "8",
         "Akt2, Slc2a4, Gsk3b, Akt1, Irs1, Irs2, Insr, Foxo1"),
        ("PPAR_Signaling",
         "Investigator-curated (KEGG mmu03320)",
         "8",
         "Lpl, Fabp4, Fabp3, Pparg, Ppara, Rxra, Acox1, Scd1"),
        ("Cardiac_Fibrosis",
         "Investigator-curated (published literature)",
         "9",
         "Acta2, Vim, Col1a1, Col1a2, Col3a1, Tgfb1, Postn, Fn1, Ctgf"),
    ]

    for row_num, (gs_name, source, n, genes) in enumerate(data_rows, start=1):
        row   = table.add_row()
        cells = row.cells
        values = [gs_name, source, n, genes]
        for idx, cell in enumerate(cells):
            p_cell = cell.paragraphs[0]
            set_para_spacing(p_cell, space_before=30, space_after=30)
            add_run(p_cell, values[idx], bold=False, font_size=9, color=BLACK)

    set_column_widths(table, col_widths_s2)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Building document ...")
    doc = create_document()

    print("  Adding Table S1 (Key Resources Table) ...")
    add_table_s1(doc)

    print("  Adding Table S2 (GSEA Gene Set Definitions) ...")
    add_table_s2(doc)

    # Ensure output directory exists
    out_dir = os.path.dirname(OUTPUT_PATH)
    os.makedirs(out_dir, exist_ok=True)

    doc.save(OUTPUT_PATH)

    size_bytes = os.path.getsize(OUTPUT_PATH)
    size_kb    = size_bytes / 1024
    print(f"\nSaved: {OUTPUT_PATH}")
    print(f"File size: {size_bytes:,} bytes ({size_kb:.1f} KB)")
    print("Done.")


if __name__ == "__main__":
    main()
