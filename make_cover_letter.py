"""
make_cover_letter.py
Generates iScience cover letter DOCX for HFpEF multi-omics manuscript
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# ── Paths (relative to repo root) ────────────────────────────────────────────
import os as _os
BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))
_os.makedirs(_os.path.join(BASE_DIR, 'figures_out'), exist_ok=True)

PUB  = _os.path.join(BASE_DIR, 'figures_out')
DATA = _os.path.join(BASE_DIR, 'data')
DEG  = _os.path.join(BASE_DIR, 'deg')


OUT = _os.path.join(BASE_DIR, 'submission_package/CoverLetter/CoverLetter_iScience_HFpEF.docx')

doc = Document()
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
    setattr(sec, attr, Inches(1.0))

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

def para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    return p

def blank():
    para('')

# ── Date & sender ────────────────────────────────────────────────────────────
today = datetime.date.today().strftime('%B %d, %Y')
para(today, space_after=12)

para('Liangqing Zhang, PhD')
para('Guangdong Medical University')
para('Zhanjiang, Guangdong, China')
para('zhangliangqing@gdmu.edu.cn', space_after=16)

# ── Addressee ────────────────────────────────────────────────────────────────
para('The Editors')
para('iScience', bold=True)
para('Cell Press / Elsevier', space_after=16)

# ── Salutation ───────────────────────────────────────────────────────────────
para('Dear Editors,', space_after=10)

# ── Opening paragraph ────────────────────────────────────────────────────────
para(
    'We are pleased to submit our manuscript entitled '
    '"A Multi-Omics Framework Reveals Biphasic Metabolic Dysregulation and a '
    'PPAR\u03b1\u2013miR-29b\u2013TGF-\u03b21 Axis in Heart Failure with Preserved '
    'Ejection Fraction" for consideration as an Article in iScience. '
    'This work has not been published previously, is not under consideration '
    'elsewhere, and has been approved by all co-authors.',
    space_after=10
)

# ── Scientific rationale ─────────────────────────────────────────────────────
para(
    'Heart failure with preserved ejection fraction (HFpEF) now accounts for '
    'more than half of all heart failure hospitalizations, yet its molecular '
    'mechanisms remain insufficiently characterized and no disease-modifying '
    'therapy has demonstrated clear benefit beyond SGLT2 inhibitors. A central '
    'unresolved question is how the cardiac metabolic phenotype evolves across '
    'HFpEF progression, and how metabolic dysfunction couples to fibrotic '
    'remodeling to drive diastolic dysfunction.',
    space_after=10
)

# ── Key findings ─────────────────────────────────────────────────────────────
para(
    'Using integrative multi-omics analysis spanning three independent publicly '
    'available datasets — mouse early-phase HFpEF cardiomyocytes (GSE209548), '
    'established mouse HFpEF whole-heart transcriptomics (GSE194151; Cao et al., '
    'Nat. Commun. 2022), and human plasma miRNA profiling (GSE53437) — and four '
    'complementary analytical layers (pre-ranked GSEA, STRING v12.0 protein '
    'interaction network, PPARα structure-based virtual screening, and plasma '
    'miRNA differential expression), we report the following key findings:',
    space_after=6
)

findings = [
    ('Proposed biphasic metabolic reprogramming: ',
     'Early-phase HFpEF cardiomyocytes exhibit coordinated pan-metabolic '
     'suppression (FAO NES=−1.90, glycolysis NES=−1.84, insulin signaling '
     'NES=−1.71; all FDR<0.05) without fibrotic activation — a pre-fibrotic '
     'metabolic state. Established HFpEF transitions to compensatory oxidative '
     'phosphorylation upregulation (NES=+1.87, FDR=0.007) with persistent '
     'PPARα/FAO suppression and induction of individual fibrosis effectors '
     '(Col1a1, Col3a1, Tgfb1). This framework may reconcile previously '
     'conflicting reports of metabolic gene expression in HFpEF.'),
    ('Network topology identifies fibrotic hubs: ',
     'STRING v12.0 network analysis (23 seed genes; expanded to 856 proteins, '
     '1,617 edges at confidence ≥0.700) positions fibronectin-1 (Fn1; '
     'degree=227) and TGF-β1 (degree=167) as the primary protein interaction '
     'hubs linking metabolic dysfunction to ECM remodeling.'),
    ('Computational prioritization of pemafibrate: ',
     'AutoDock Vina molecular docking against the PPARα ligand-binding domain '
     '(PDB:1K7L) identifies pemafibrate (ΔG=−10.5 kcal/mol) as the top '
     'translationally prioritized candidate among six profiled compounds, '
     'corroborated by literature-sourced CMAP connectivity scores '
     '(pemafibrate ranked first, +0.68).'),
    ('PPARα–miR-29b–TGF-β1 therapeutic axis: ',
     'Human plasma miRNA profiling confirms concordant depletion of '
     'hsa-miR-423-3p/5p as established HF biomarkers, and identifies a '
     'sub-threshold trend for hsa-miR-29b-3p depletion (padj=0.064), '
     'whose ECM target specificity (COL1A1 miRDB=94, COL3A1=91) suggests '
     'a candidate mechanistic link to the fibrotic program. Together, '
     'these data support a proposed PPARα–miR-29b–TGF-β1 multi-target axis '
     'as a framework for HFpEF intervention, pending in vivo validation.'),
]

for bold_part, normal_part in findings:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f'• {bold_part}')
    r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.name = 'Arial'; r2.font.size = Pt(11)

# ── Significance ─────────────────────────────────────────────────────────────
para(
    'We believe this work is well-suited to iScience\'s multidisciplinary scope '
    'because it integrates transcriptomics, network biology, structural '
    'pharmacology, and miRNA biology within a single analytical framework '
    'applied to an important and growing clinical problem. All datasets are '
    'publicly available, and analysis code is deposited on GitHub, ensuring '
    'full reproducibility.',
    space_before=6, space_after=10
)

# ── Suggested reviewers ───────────────────────────────────────────────────────
para('We suggest the following potential reviewers (no conflicts of interest):',
     space_after=6)

reviewers = [
    ('[Reviewer 1 Name], [Institution] — expert in cardiac metabolism and HFpEF]',
     '[reviewer1@institution.edu]'),
    ('[Reviewer 2 Name], [Institution] — expert in cardiac fibrosis / TGF-β signaling]',
     '[reviewer2@institution.edu]'),
    ('[Reviewer 3 Name], [Institution] — expert in computational pharmacology / PPARα]',
     '[reviewer3@institution.edu]'),
]
for name, email in reviewers:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'• {name}  {email}')
    r.font.name = 'Arial'; r.font.size = Pt(11)

# ── Closing ──────────────────────────────────────────────────────────────────
para(
    'We confirm that this manuscript is original, has not been published, and '
    'is not currently under consideration at any other journal. All authors '
    'have read and approved the final manuscript and consent to its submission.',
    space_before=10, space_after=10
)

para('Thank you for your consideration.', space_after=16)
para('Sincerely,', space_after=24)
para('Liangqing Zhang, PhD', bold=True)
para('Guangdong Medical University')
para('Zhanjiang, Guangdong, China')
para('zhangliangqing@gdmu.edu.cn')

import os
doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size:  {os.path.getsize(OUT):,} bytes')
