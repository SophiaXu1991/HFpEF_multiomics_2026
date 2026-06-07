"""
write_manuscript_final.py
iScience (Cell Press) format – proper heading hierarchy, section order, styling
Output: iScience_HFpEF_manuscript_final.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import os

# ── Paths (relative to repo root) ────────────────────────────────────────────
import os as _os
BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))
_os.makedirs(_os.path.join(BASE_DIR, 'figures_out'), exist_ok=True)

PUB  = _os.path.join(BASE_DIR, 'figures_out')
DOCX = _os.path.join(BASE_DIR, 'iScience_HFpEF_manuscript_v28.docx')

# ── Document setup ──────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
    setattr(sec, attr, Inches(1.0))

# Base font
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

# Style Heading 1 – major sections (Introduction, Results, Discussion…)
h1 = doc.styles['Heading 1']
h1.font.name  = 'Arial'
h1.font.size  = Pt(13)
h1.font.bold  = True
h1.font.color.rgb = RGBColor(0, 70, 127)       # Cell Press dark blue
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after  = Pt(6)
h1.paragraph_format.keep_with_next = True

# Style Heading 2 – Results subsections, Methods subsections
h2 = doc.styles['Heading 2']
h2.font.name  = 'Arial'
h2.font.size  = Pt(11)
h2.font.bold  = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after  = Pt(3)
h2.paragraph_format.keep_with_next = True

# Style Heading 3 – sub-subsections (STAR Methods sub-items)
h3 = doc.styles['Heading 3']
h3.font.name  = 'Arial'
h3.font.size  = Pt(11)
h3.font.bold  = False
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after  = Pt(2)
h3.paragraph_format.keep_with_next = True

# ── Helpers ─────────────────────────────────────────────────────────────────────

def h1_(text):
    doc.add_heading(text, level=1)

def h2_(text):
    doc.add_heading(text, level=2)

def h3_(text):
    doc.add_heading(text, level=3)

def body(text, sa=6, sb=0, italic=False, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.size  = Pt(11)
    run.font.italic = italic
    return p

def body_mixed(parts, sa=6, sb=0, justify=True):
    """parts: list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name   = 'Arial'
        r.font.size   = Pt(11)
        r.bold        = bold
        r.font.italic = italic
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)

def embed_fig(filename, width=6.2):
    path = os.path.join(PUB, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        ph = doc.add_paragraph(f'[Figure not found: {filename}]')
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fig_caption(label, text, sa=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(9)

def fig_legend_entry(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)

def ref_entry(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after         = Pt(3)
    p.paragraph_format.first_line_indent   = Inches(-0.3)
    p.paragraph_format.left_indent         = Inches(0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f'{num}.\u2002{text}')
    r.font.name = 'Arial'
    r.font.size = Pt(10)

def page_break():
    doc.add_page_break()

def thin_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(8)
tr = title_p.add_run(
    'A Multi-Omics Framework Reveals Biphasic Metabolic Dysregulation and a '
    'PPAR\u03b1\u2013miR-29b\u2013TGF-\u03b21 Axis in Heart Failure with Preserved Ejection Fraction'
)
tr.bold = True
tr.font.size = Pt(15)
tr.font.name = 'Arial'

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(3)
ar = auth_p.add_run(
    'Riping Xu\u1d43\u2020, Li Huang\u1d43\u2020, Tao Chen\u1d47\u2020, Liangqing Zhang\u1d9c\u002a'
)
ar.font.size = Pt(10); ar.font.name = 'Arial'

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(2)
af = aff_p.add_run(
    '\u1d43The Department of Anesthesiology, Affiliated Hospital of Guangdong Medical University, '
    'Zhanjiang, Guangdong, China; '
    '\u1d47Department of Cardiology, The Affiliated Hospital of Guangdong Medical University, '
    'Zhanjiang, Guangdong, China; '
    '\u1d9cGuangdong Medical University, Zhanjiang, Guangdong, China'
)
af.font.size = Pt(9); af.font.name = 'Arial'; af.font.italic = True

eq_p = doc.add_paragraph()
eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_p.paragraph_format.space_after = Pt(3)
eq = eq_p.add_run(
    '\u2020Riping Xu, Li Huang, and Tao Chen contributed equally to this work.'
)
eq.font.size = Pt(9); eq.font.name = 'Arial'; eq.font.italic = True

corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_p.paragraph_format.space_after = Pt(12)
cr = corr_p.add_run(
    '\u002aCorrespondence: Liangqing Zhang, PhD  \u2014  zhangliangqing@gdmu.edu.cn'
)
cr.font.size = Pt(10); cr.font.name = 'Arial'; cr.font.italic = True

thin_rule()

# ══════════════════════════════════════════════════════════════════════════════
# HIGHLIGHTS
# ══════════════════════════════════════════════════════════════════════════════

hl_p = doc.add_paragraph()
hl_p.paragraph_format.space_before = Pt(8)
hl_p.paragraph_format.space_after  = Pt(4)
hlr = hl_p.add_run('Highlights')
hlr.bold = True; hlr.font.size = Pt(12); hlr.font.name = 'Arial'
hlr.font.color.rgb = RGBColor(0, 70, 127)

bullet(
    'Early HFpEF cardiomyocytes show pan-metabolic suppression without fibrotic activation'
)
bullet(
    'Established HFpEF combines OxPhos upregulation with persistent PPAR\u03b1/FAO suppression'
)
bullet(
    'STRING network places Fn1 (degree=227) and TGF-\u03b21 (167) as top fibrotic hub nodes'
)
bullet(
    'PPAR\u03b1 virtual screening nominates pemafibrate (\u0394G=\u221210.5\u2009kcal/mol) as top computational candidate'
)
bullet(
    'Plasma miR-29b-3p trend and CMAP scores support a proposed PPAR\u03b1\u2013miR-29b\u2013TGF-\u03b21 axis'
)

thin_rule()

# ══════════════════════════════════════════════════════════════════════════════
# IN BRIEF  (eTOC blurb — required by iScience)
# ══════════════════════════════════════════════════════════════════════════════

ib_p = doc.add_paragraph()
ib_p.paragraph_format.space_before = Pt(6)
ib_p.paragraph_format.space_after  = Pt(2)
ibr = ib_p.add_run('In Brief')
ibr.bold = True; ibr.font.size = Pt(11); ibr.font.name = 'Arial'
ibr.font.color.rgb = RGBColor(0, 70, 127)

body(
    'Integrative multi-omics analysis of heart failure with preserved ejection fraction '
    '(HFpEF) across three independent datasets identifies a proposed biphasic metabolic '
    'reprogramming — pan-metabolic suppression in early disease transitioning to '
    'compensatory oxidative phosphorylation upregulation in established HFpEF. '
    'STRING network topology, PPARα virtual screening, and human plasma miRNA profiling '
    'converge on a PPARα–miR-29b–TGF-β1 multi-target axis as a candidate therapeutic '
    'framework, with pemafibrate nominated as the top computationally prioritized agent '
    'pending in vivo validation.'
)

thin_rule()

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY  (iScience calls it "Summary", not "Abstract")
# ══════════════════════════════════════════════════════════════════════════════

h1_sum = doc.add_paragraph()
h1_sum.paragraph_format.space_before = Pt(8)
h1_sum.paragraph_format.space_after  = Pt(4)
h1r = h1_sum.add_run('Summary')
h1r.bold = True; h1r.font.size = Pt(13); h1r.font.name = 'Arial'
h1r.font.color.rgb = RGBColor(0, 70, 127)

body(
    'Heart failure with preserved ejection fraction (HFpEF) now accounts for more than half '
    'of all heart failure hospitalizations and carries five-year mortality approaching 50%, '
    'yet its molecular mechanisms remain insufficiently characterized. We performed '
    'integrative multi-omics analysis across three independent datasets — mouse early-phase '
    'HFpEF cardiomyocytes (GSE209548; wild-type high-fat diet vs. chow control, n=3/3), '
    'established mouse HFpEF whole-heart transcriptomics (GSE194151; '
    'HFD+L-NAME model, n=15 HFpEF vs. n=15 control, sex-pooled; Cao et al., 2022), '
    'and human plasma miRNA profiling '
    '(GSE53437, n=29/14) — to characterize molecular events driving diastolic '
    'dysfunction. Gene set enrichment analysis identified a proposed biphasic pattern of '
    'metabolic reprogramming: early coordinated suppression of fatty acid oxidation '
    '(FAO; NES=−1.90, FDR=0.019), glycolysis (NES=−1.84, FDR=0.016), and insulin '
    'signaling (NES=−1.71, FDR=0.040) transitions in established disease to compensatory '
    'oxidative phosphorylation upregulation (NES=+1.87, FDR=0.007) alongside persistent '
    'PPAR\u03b1/FAO suppression and upregulation of individual fibrosis effectors '
    '(Col1a1, Col3a1, Tgfb1) at the DEG level — without coordinate pathway-level '
    'GSEA enrichment of the Cardiac_Fibrosis gene set. STRING protein interaction network '
    'analysis identified fibronectin-1 (Fn1; degree=227) and TGF-β1 (degree=167) as primary '
    'hub nodes mechanistically linking metabolic dysfunction to ECM remodeling. PPARα virtual '
    'screening identified GW7647 (ΔG=−11.2 kcal/mol) and pemafibrate (ΔG=−10.5 kcal/mol) '
    'as high-affinity candidates; human plasma miRNA profiling revealed concordant '
    'downregulation of hsa-miR-423-3p/5p and a sub-threshold depletion trend for '
    'hsa-miR-29b-3p (padj=0.064, not conventionally significant), whose high-confidence '
    'ECM targets (COL1A1 miRDB=94, COL3A1=91) suggest a putative mechanistic link between '
    'plasma miRNA loss and fibrotic activation pending experimental confirmation. '
    'Literature-sourced CMAP connectivity scores ranked pemafibrate as the top '
    'transcriptional reversal agent (connectivity +0.68), supporting a proposed '
    'PPARα–miR-29b–TGF-β1 putative regulatory axis as a computational hypothesis for HFpEF pending experimental validation.'
)

body_mixed([
    ('Keywords: ', True, False),
    ('HFpEF; biphasic metabolic reprogramming; fatty acid oxidation; PPARα; cardiac '
     'fibrosis; TGF-β1; miR-29b-3p; STRING network; molecular docking; pemafibrate; '
     'GSEA; multi-omics', False, False),
], sa=8)

thin_rule()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1_('Introduction')

body(
    'Heart failure with preserved ejection fraction (HFpEF) is the dominant and '
    'fastest-growing form of heart failure, accounting for more than half of all '
    'incident cases and carrying five-year mortality exceeding 50% (McDonagh et al., 2021; '
    'Seferovic et al., 2019; Borlaug, 2014; Borlaug and Paulus, 2011). Its clinical '
    'hallmarks — diastolic dysfunction with preserved left ventricular ejection fraction, '
    'exercise intolerance, and elevated filling pressures — arise from a complex interplay '
    'of myocardial fibrosis, microvascular inflammation, metabolic cardiomyopathy, and '
    'hemodynamic stress, frequently compounded by obesity and hypertension. Despite this '
    'burden, pharmacological progress has been frustratingly slow: SGLT2 inhibitors '
    '(empagliflozin and dapagliflozin) have recently demonstrated modest but significant '
    'reductions in hospitalization, providing the first disease-modifying signals in this '
    'population (Anker et al., 2021; Solomon et al., 2022). However, the molecular '
    'mechanisms underlying HFpEF pathogenesis remain insufficiently characterized to '
    'rationalize combination therapeutic strategies.'
)
body(
    'The adult heart is the highest-energy-consuming organ in the body, deriving 60–70% '
    'of its ATP from mitochondrial fatty acid oxidation (FAO) under physiological conditions '
    '(Stanley et al., 2005; Lopaschuk et al., 2010; Kolwicz et al., 2013). PPARα '
    '(peroxisome proliferator-activated receptor alpha) serves as the master transcriptional '
    'regulator of cardiac FAO, controlling rate-limiting enzymes including CPT1B, HADHA, '
    'and ACADM. In heart failure, progressive metabolic inflexibility — the loss of capacity '
    'to dynamically switch between FAO and glucose oxidation — is a defining '
    'pathophysiological feature that precedes and amplifies contractile dysfunction '
    '(Karwi et al., 2018; Lopaschuk et al., 2021). While FAO suppression is well-documented '
    'in heart failure with reduced ejection fraction (HFrEF), the temporal sequence of '
    'substrate metabolism remodeling across HFpEF progression has not been systematically '
    'characterized (Kolwicz et al., 2013; Karwi et al., 2018).'
)
body(
    'Concurrent with metabolic remodeling, cardiac fibrosis driven by TGF-β1 signaling is '
    'a pathological hallmark of HFpEF diastolic stiffness (Shah and Solomon, 2012). ECM '
    'remodeling through SMAD2/3 phosphorylation and myofibroblast differentiation deposits '
    'collagen I/III and fibronectin, increasing passive myocardial stiffness and impairing '
    'lusitropy. A largely unexplored dimension bridging metabolic and fibrotic axes is the '
    'circulating microRNA (miRNA) landscape. Circulating miRNAs have emerged as both '
    'diagnostic biomarkers and mechanistic effectors in cardiovascular disease '
    '(Creemers et al., 2012). hsa-miR-29b-3p, a post-transcriptional repressor of ECM '
    'genes (COL1A1, COL3A1, FBN1), is reduced in cardiac fibrosis contexts '
    '(Van Rooij et al., 2008), yet its role in the HFpEF plasma miRNome has not been '
    'systematically characterized.'
)
body(
    'Here we address this knowledge gap through integrative multi-omics analysis spanning '
    'three independent datasets, two species, and four complementary analytical layers: '
    '(1) cross-temporal transcriptomic GSEA to characterize a proposed biphasic metabolic reprogramming pattern; '
    '(2) STRING v12.0 protein interaction network analysis to identify hub nodes connecting '
    'metabolic and fibrotic axes; (3) PPARα structure-based virtual screening to prioritize '
    'therapeutic candidates; and (4) human plasma miRNA differential expression to link '
    'circulating biomarkers to fibrotic mechanisms. Our analyses converge on a proposed '
    'PPARα\u2013miR-29b-3p\u2013TGF-β1 putative regulatory axis as a computational hypothesis '
    'for HFpEF investigation, with pemafibrate emerging as the top computationally '
    'prioritized candidate from virtual screening and literature-sourced CMAP connectivity scores.'
)

# ══════════════════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════════════════

h1_('Results')

body(
    'The integrated multi-omics analysis workflow, spanning three independent datasets '
    'and four analytical layers, is summarized in Figure 1.'
)

doc.add_paragraph()
embed_fig('Fig1_final.png', width=6.3)
fig_caption(
    'Figure 1.',
    'Multi-Omics Study Design and Integrated Analytical Framework. Overview of three '
    'independent datasets (GSE209548, GSE194151, GSE53437) and the four-layer analytical '
    'pipeline: DESeq2 differential expression, fgsea pathway enrichment, STRING v12.0 '
    'protein interaction network, PPARα virtual screening, and literature-sourced '
    'CMAP/L1000 drug repurposing connectivity scores. '
    'Findings converge on a proposed PPARα–miR-29b–TGF-β1 fibrotic signaling axis.'
)

# ── Results 2.1 ──────────────────────────────────────────────────────────────
h2_('Phase 1 (HFD Model, GSE209548) — Early Metabolic Stress: Coordinated '
    'Pan-Metabolic Suppression Without Fibrotic Activation')

body(
    'To characterize early-phase molecular events, we analyzed GSE209548, '
    'a mouse cardiomyocyte-enriched RNA-seq dataset from the IRCCS Humanitas Research '
    'Hospital (Mus musculus). The dataset contains four groups (wild-type [WT] chow, '
    'WT high-fat diet [HFD], ApoE KO chow, ApoE KO HFD; n=3/group; 12 samples total); '
    'this analysis used the WT arm only (ApoE KO samples excluded), comparing '
    'WT HFD vs. WT chow control (n=3 vs. n=3). This design isolates the dietary '
    'high-fat effect in a normal genetic background, without the confounding dyslipidemia '
    'of the ApoE KO genotype. DESeq2 differential expression analysis '
    '(FDR<0.05, |log2FC|>0.5) did not identify individually significant DEGs at this '
    'early time point, consistent with the subtle transcriptional perturbation '
    'characteristic of metabolic stress prior to established cardiomyopathy. This absence '
    'of discrete gene-level changes motivates a pathway-level analytical approach.'
)
body(
    'Pre-ranked GSEA revealed a striking pattern of coordinated pathway-level suppression '
    '(Figure S1). Fatty acid oxidation (FAO_KEGG) showed the strongest suppression '
    '(NES=−1.90, FDR=0.019), followed by glycolysis (NES=−1.84, FDR=0.016) and insulin '
    'signaling (NES=−1.71, FDR=0.040). OxPhos exhibited a non-significant suppression '
    'trend (NES=−1.58, FDR=0.078), while cardiac fibrosis showed no enrichment '
    '(NES=+0.99, FDR=0.513). This dissociation between metabolic suppression and fibrotic '
    'quiescence suggests Phase 1 as a pre-fibrotic state of coordinated pan-metabolic '
    'energy substrate suppression, consistent with the hypothesis that metabolic '
    'reprogramming is an early event in HFpEF pathogenesis.'
)

body(
    'Cross-phase pathway-level comparison suggests a directional reversal of '
    'metabolic enrichment between Phase 1 and Phase 2: FAO and glycolysis shift from '
    'suppressed (NES<\u22121.5) in Phase 1 to non-significantly suppressed in Phase 2 '
    '(FAO_KEGG NES=\u22121.21, FDR=0.300), while OxPhos transitions from non-significant '
    'suppression (NES=\u22121.58, FDR=0.078) to significant activation '
    '(NES=+1.87, FDR=0.007). The Cardiac_Fibrosis pathway gene set remains '
    'non-significant in both phases (Phase 1 NES=+0.99, FDR=0.513; Phase 2 '
    'NES=\u22120.99, FDR=0.519); active fibrotic remodeling in Phase 2 is supported '
    'by individual DEG upregulation rather than pathway-level GSEA enrichment (Figure 2).'
)

doc.add_paragraph()
embed_fig('FigA_biphasic_NES.png', width=6.2)
fig_caption(
    'Figure 2.',
    'Proposed Biphasic Metabolic Reprogramming: Pathway-Level NES Comparison Across '
    'Two Independent Mouse HFpEF Models. Grouped bar chart showing normalized enrichment '
    'scores (NES) for seven GSEA pathways (FAO, Glycolysis, Insulin Signaling, OxPhos, '
    'PPAR\u03b1 Signaling, Ketone Metabolism, Cardiac Fibrosis) in the HFD model '
    '(Phase 1, GSE209548; WT high-fat diet vs. chow control, n=3/3; red) and the '
    'HFD+L-NAME model (Phase 2, GSE194151; whole heart, n=15 vs. n=15, sex-pooled; '
    'Cao et al., 2022; blue). Note: the two models differ in stressor, tissue, and mouse '
    'background; their comparison infers rather than directly demonstrates temporal '
    'progression. Saturated color: FDR<0.10; faded color: FDR\u22650.10 (not significant). '
    'Significance: ** FDR<0.05, # FDR 0.05\u20130.10 (trend), ns FDR\u22650.10. '
    'Arrow indicates OxPhos transition from suppression (Phase 1) to activation (Phase 2). '
    'Dashed lines: NES=\u00b11.5 reference threshold.'
)

# ── Results 2.2 ──────────────────────────────────────────────────────────────
h2_('Phase 2 (HFD+L-NAME Model, GSE194151) — Established HFpEF: Compensatory '
    'OxPhos Upregulation with Persistent FAO Depression')

body(
    'To characterize the established HFpEF transcriptomic landscape, we analyzed '
    'GSE194151 (Cao et al., 2022), a mouse whole-heart RNA-seq dataset from the '
    'HFD+L-NAME two-hit model (C57BL/6J; Lusis laboratory, UCLA). The original study '
    'was designed to investigate sex differences in cardiac mitochondria and diastolic '
    'dysfunction (15 male + 15 female mice); for this analysis, males and females were '
    'pooled to form HFD+L-NAME vs. control comparison groups (n=15/15 after pooling), '
    'with sex not included as a covariate. Pooling sexes increases power but introduces '
    'sex as an uncontrolled variable; results should be interpreted with this caveat. '
    'Importantly, GSE194151 also uses distinct whole-heart tissue compared with the '
    'cardiomyocyte-enriched GSE209548 dataset, and a different mouse background and '
    'metabolic stress mechanism (L-NAME + HFD vs. dietary HFD alone in WT mice); '
    'the cross-dataset comparison therefore represents an inference of disease '
    'progression rather than direct longitudinal tracking within a single model. '
    'DESeq2 differential expression analysis (Love et al., 2014; Muzellec et al., 2023) '
    'identified 7,921 significant DEGs: 3,084 upregulated and 4,837 downregulated. Key '
    'downregulated FAO genes included Cpt1b (log2FC=−3.4), Hadha (log2FC=−2.8), '
    'Acadm (log2FC=−2.1), Hadhb (log2FC=−2.2), and the master regulator Ppara itself '
    '(log2FC=−1.9), confirming persistent and pervasive PPARα/FAO pathway suppression '
    'at the gene-expression level (Figure 3A).'
)
body(
    'GSEA running score analysis revealed a paradoxical transcriptional landscape '
    '(Figure 3B). OxPhos genes were significantly upregulated in established HFpEF '
    '(OxPhos_KEGG: NES=+1.87, FDR=0.007), indicating transcriptional upregulation of '
    'electron transport chain (ETC) components as a compensatory response to sustained '
    'energetic stress. FAO pathway genes remained suppressed (FAO_KEGG: NES=\u22121.21, '
    'FDR=0.300). At the individual gene level, key cardiac fibrosis effectors '
    '(Col1a1, Col3a1, Postn, Fn1, Tgfb1, Acta2) were significantly upregulated in the '
    'DEG analysis (Figure 3A), consistent with active ECM remodeling; however, the '
    'Cardiac_Fibrosis pathway gene set did not reach significance in the pre-ranked GSEA '
    '(NES=\u22120.99, FDR=0.519), likely reflecting heterogeneous pathway gene-set composition '
    'rather than the absence of fibrotic activation. The co-occurrence of OxPhos '
    'upregulation and FAO depression \u2014 representing Phase 2 of the biphasic model \u2014 '
    'is consistent with, though does not directly prove, a substrate switch toward '
    'alternative oxidative fuels (amino acids, ketone bodies) under conditions of '
    'impaired FAO flux; isotope-tracing or metabolomic data would be required to '
    'confirm this inference.'
)
body(
    'WGCNA co-expression network analysis (Langfelder and Horvath, 2008) of GSE194151 '
    '(n=30 samples) identified a Blue module as the FAO co-expression core, with '
    'highest-membership hub genes Cpt1b (kME=0.951), Oxct1 (kME=0.963), Smyd1 '
    '(kME=0.967), and Sspn (kME=0.971). The high module membership of Smyd1 — a '
    'histone methyltransferase known to regulate metabolic gene programs in '
    'cardiomyocytes — and Sspn (sarcospan, a structural membrane scaffolding protein '
    'co-regulated with cytoskeletal remodeling) likely reflects broad transcriptional '
    'co-regulation within the FAO-depressed state rather than direct enzymatic FAO '
    'function. The Blue module was inversely correlated with the Turquoise module '
    '(r=−0.960, p<0.001; Figure S3), quantifying the metabolic-fibrotic axis '
    'bifurcation in the established HFpEF state (GSE194151).'
)

body(
    'To visualize per-sample expression patterns of the key metabolic and fibrotic '
    'gene sets in established HFpEF, we generated a z-scored log2-CPM heatmap across '
    'all 30 samples (Figure 3C). FAO/PPARα genes (red) were uniformly downregulated '
    'in HFpEF relative to controls, OxPhos genes (blue) showed upregulation in most '
    'HFpEF samples consistent with compensatory ETC transcription, and ECM/Fibrosis '
    'genes (green) were markedly elevated, providing direct per-sample visual '
    'confirmation of the biphasic transcriptional program.'
)

doc.add_paragraph()
embed_fig('Fig_merged3.png', width=6.4)
fig_caption(
    'Figure 3.',
    'Transcriptomic Analysis of Established HFpEF (GSE194151; HFD+L-NAME vs. control; '
    'n=15/15 sex-pooled; Cao et al., 2022). '
    '(A) Volcano plot: red=upregulated (n=3,084; padj<0.05, |log2FC|>0.5), '
    'blue=downregulated (n=4,837); key FAO, OxPhos, and fibrosis genes labeled. '
    '(B) GSEA NES bar chart for seven pathway gene sets; OxPhos_KEGG enriched '
    '(NES=+1.87, FDR=0.007); FAO_KEGG suppressed; ** FDR<0.05, # FDR 0.05\u20130.10 (trend). '
    '(C) Per-sample z-scored log2-CPM heatmap for 27 FAO/PPARα (red), OxPhos (blue), '
    'and ECM/Fibrosis (green) genes across 15 Control and 15 HFpEF samples.'
)

# ── Results 2.3 (robustness) ─────────────────────────────────────────────────
h2_('Sex-Stratified Robustness and Cross-Model Validation of Phase 2 Metabolic Signature')

body(
    'To address two key concerns about the Phase 2 GSEA findings — potential sex-specific '
    'confounding in the pooled GSE194151 analysis and L-NAME pharmacological effects on '
    'OxPhos gene programs — we conducted two additional analyses.'
)
body(
    'Sex-stratified GSEA: GSE194151 contains 15 male and 15 female mice (identified via '
    '"Sex: M/F" GEO metadata). We re-ran DESeq2 and pre-ranked GSEA independently for '
    'male (HFpEF n=8, control n=8) and female (HFpEF n=7, control n=7) subsets. '
    'All seven pathway NES values were directionally concordant between sexes '
    '(Pearson r=0.996, 7/7 concordant; Figure S4). Critically, OxPhos upregulation was '
    'independently significant in both males (NES=+1.73, FDR=0.005) and females '
    '(NES=+1.45, FDR=0.032), and ketone suppression reached significance in males '
    '(NES=\u22121.64, FDR=0.006) with a consistent trend in females (NES=\u22121.23, FDR=0.204). '
    'PPAR\u03b1 signaling suppression trended significantly in males (NES=\u22121.38, FDR=0.085). '
    'These results confirm that the metabolic signature of Phase 2 is not attributable '
    'to sex imbalance and is present in both sexes independently.'
)
body(
    'Cross-model validation (L-NAME-free dataset): To evaluate whether the Phase 2 OxPhos '
    'upregulation could be confounded by L-NAME\'s direct effects on mitochondrial gene programs, '
    'we analyzed GSE249409 — an independent HFpEF dataset using the HFD+pressure-overload '
    '(mTAC) model without L-NAME (n=5 control, n=5 HFpEF). In this L-NAME-free model, '
    'FAO suppression was maintained (NES=\u22120.81) and PPAR\u03b1 signaling showed a positive '
    'but non-significant trend (NES=+1.02, FDR=0.95), consistent with partial metabolic '
    'overlap across HFpEF models. However, OxPhos enrichment was absent and directionally '
    'reversed (NES=\u22120.77, FDR=0.76), suggesting that the robust OxPhos upregulation '
    'observed in GSE194151 may reflect synergistic effects of HFD and NOS inhibition '
    'specific to the HFD+L-NAME model, rather than a universal Phase 2 feature. '
    'We note that all GSE249409 NES values were non-significant (all FDR>0.60), '
    'likely reflecting the limited power of n=5/5. This finding is discussed as a '
    'limitation in the context of the biphasic model below.'
)

doc.add_paragraph()
embed_fig('FigureS4_SexStratified_GSEA.png', width=6.4)
fig_caption(
    'Figure S4.',
    'Sex-Stratified GSEA Robustness Analysis — GSE194151 (HFpEF vs. Control). '
    'Male (n=8/group, blue) and female (n=7/group, pink) mice analyzed independently '
    'via DESeq2 and pre-ranked GSEA. Bar heights show NES for seven pathway gene sets; '
    'full opacity = FDR<0.10, reduced opacity = FDR≥0.10 (non-significant). '
    '** FDR<0.05, # FDR 0.05\u20130.10 (trend), ns FDR\u22650.10. All seven pathway NES values are directionally concordant '
    'between sexes (7/7 concordant). OxPhos upregulation is independently significant '
    'in both males (NES=+1.73, FDR=0.005) and females (NES=+1.45, FDR=0.032). '
    'Dashed lines: NES=±1.5 (threshold for strong enrichment).'
)

# ── Results 2.4 (was 2.3) ────────────────────────────────────────────────────
h2_('STRING Protein Interaction Network Identifies Fibrotic Hub Genes as '
    'Primary Integrators')

body(
    'To map protein-level interactions connecting metabolic and fibrotic regulators, we '
    'queried STRING v12.0 (Szklarczyk et al., 2023) using 23 seed genes selected from '
    'GSEA-prioritized pathways: GSE194151-significant DEGs (padj<0.05, |log2FC|>0.5) '
    'from six gene set categories were filtered, then augmented by ten mandatory metabolic '
    'and fibrotic anchor genes (Cpt1b, Acadm, Ppara, Cycs, Hmgcs2, Fn1, Col1a1, Tgfb1, '
    'Postn, Acta2). The 23 unique seed genes span FAO (Cpt1b, Hadha, Acadm, Hadhb), '
    'ketone metabolism (Hmgcs2, Oxct1, Bdh1), OxPhos (Ndufa1, Cycs, Cox4i1), '
    'PPARα/insulin signaling (Ppara, Foxo1, Akt1, Insr), and ECM/fibrotic '
    'effectors (Col1a1, Col1a2, Col3a1, Fn1, Tgfb1, Acta2, Postn). '
    'The STRING interaction_partners API endpoint (Mus musculus, taxon 10090, combined '
    'confidence score ≥0.700) expanded these 23 seed proteins to a 856-protein '
    'functional interaction network comprising 1,617 edges (Figure 4A). The interaction '
    'score distribution (Figure 4B) was markedly right-skewed: 314 edges (19.4%) at '
    'very high confidence (≥0.95) and 405 edges (25.0%) at high confidence (0.90–0.95); '
    'all hub-gene degree values reported below refer to this 856-protein expanded network.'
)
body(
    'Hub gene analysis by network degree (Figure 4C) identified fibronectin-1 '
    '(Fn1; degree=227) and TGF-β1 (Tgfb1; degree=167) as the two highest-connectivity '
    'nodes, substantially exceeding metabolic regulators Cox4i1 (degree=133) and '
    'Foxo1 (115). The dominance of fibrotic mediators as network hubs — surpassing the '
    'master metabolic regulator Ppara (degree=108) — suggests that ECM remodeling acts '
    'as the primary protein interaction integrator of HFpEF pathological signals, '
    'amplifying downstream collagen deposition and diastolic stiffness.'
)

doc.add_paragraph()
embed_fig('Fig3_STRING_combined.png', width=6.2)
fig_caption(
    'Figure 4.',
    'STRING v12.0 Protein Interaction Network Analysis of Key HFpEF Genes. '
    '(A) Expanded interaction network (Mus musculus, confidence ≥0.700): 23 seed genes '
    'expanded via interaction_partners API to 856 proteins, 1,617 edges. '
    '(B) Interaction score distribution; 314 edges at very high confidence (≥0.95). '
    '(C) Hub gene connectivity (degree) for top 12 genes; Fn1 (degree=227) and '
    'Tgfb1 (167) identified as primary fibrotic hub nodes.'
)

# ── Results 2.4 ──────────────────────────────────────────────────────────────
h2_('PPARα Virtual Screening Identifies High-Affinity Therapeutic Candidates')

body(
    'Given consistent PPARα suppression across both HFpEF phases, we profiled the '
    'binding affinities of six PPAR\u03b1-relevant compounds (four selective PPAR\u03b1 '
    'agonists/modulators plus two cross-PPAR comparators) against the PPAR\u03b1 '
    'ligand-binding domain (PDB:1K7L; Xu et al., 2001), using AutoDock Vina molecular '
    'docking performed in this study to compute predicted binding free energies. It should be noted that '
    'PDB:1K7L was co-crystallized with GW409544 and a corepressor peptide; its use as '
    'a docking template for agonist candidates represents an approximation that requires '
    'validation against an agonist-bound structure. Within this framework, GW7647 '
    'exhibited the strongest predicted binding affinity (ΔG=−11.2 kcal/mol), exceeding '
    'the high-affinity threshold of −10.0 kcal/mol. Pemafibrate — a selective PPARα '
    'modulator (SPPARMα) with proven triglyceride-lowering efficacy (Fruchart, 2017); '
    'the PROMINENT randomized trial confirmed its safety but did not demonstrate '
    'a reduction in major adverse cardiovascular events vs. placebo '
    '(HR=1.03, p=0.67; das Pradhan et al., 2022) — demonstrated the second-highest '
    'predicted affinity '
    '(\u0394G=\u221210.5 kcal/mol). Although GW7647 demonstrated the highest computed binding '
    'affinity, pemafibrate was designated as the translational '
    'priority candidate: GW7647 is a preclinical research tool compound without clinical '
    'development history, whereas pemafibrate holds regulatory approval in multiple '
    'jurisdictions and has an established clinical safety profile (Fruchart, 2017; '
    'das Pradhan et al., 2022). Pemafibrate therefore represents the most translationally '
    'feasible lead among the profiled compounds, pending in vivo validation '
    '(Figure 5A). WY-14643 (ΔG=−9.8 kcal/mol; Willson et al., 2000) and fenofibrate '
    '(ΔG=−9.1 kcal/mol; Staels et al., 2008) exceeded the moderate threshold '
    '(−8.0 kcal/mol).'
)
body(
    'Binding pocket interaction analysis of the GW7647–PPARα complex (Figure 5B) revealed '
    'hydrogen bond contacts with four canonical activation-function-2 (AF-2) helix '
    'residues: Tyr314, His440, Tyr464, and Ser280. Hydrophobic stabilization was provided '
    'by Leu254, Ile272, Leu347, and Phe351, consistent with interactions documented in '
    'the 1K7L crystal structure (Xu et al., 2001). Literature-sourced CMAP/L1000 '
    'connectivity scores (Subramanian et al., 2017) ranked pemafibrate as the strongest '
    'transcriptional reversal agent (+0.68), consistent with its virtual screening rank, '
    'providing dual computational evidence for its candidacy in HFpEF.'
)

doc.add_paragraph()
embed_fig('Fig4_docking.png', width=6.2)
fig_caption(
    'Figure 5.',
    'PPARα Virtual Screening and Binding Pocket Interactions. '
    '(A) AutoDock Vina binding free energies for six PPAR\u03b1-relevant compounds '
    '(GW7647, pemafibrate, WY-14643, fenofibrate; pioglitazone and GW501516 as '
    'cross-PPAR comparators; kcal/mol); '
    'dashed red line: high-affinity threshold (−10.0 kcal/mol). '
    '(B) Key binding pocket interactions in the GW7647–PPARα complex (PDB:1K7L): '
    'H-bond contacts (Tyr314, His440, Tyr464, Ser280) and hydrophobic contacts '
    '(Leu254, Ile272, Leu347, Phe351). '
    '(C–E) 2D molecular structures of GW7647 (ΔG=−11.2), pemafibrate (−10.5), '
    'and WY-14643 (−9.8 kcal/mol).'
)

# ── Results 2.5 ──────────────────────────────────────────────────────────────
h2_('Human Plasma miRNA Profiling Identifies Concordant Depletion of '
    'miR-423 and miR-29b Clusters')

body(
    'Analysis of GSE53437 (HFpEF n=29 vs. healthy n=14) identified 40 upregulated and '
    '38 downregulated miRNAs (padj<0.05; Figure S2). hsa-miR-423-3p '
    '(padj=0.022, log2FC=\u22121.63) and hsa-miR-423-5p (padj=0.036, log2FC=\u22121.49) were '
    'concordantly and significantly depleted in HFpEF plasma \u2014 established heart '
    'failure-associated circulating biomarkers first reported by Tijsen et al. (2010). '
    'hsa-miR-29b-3p showed a sub-threshold depletion trend (padj=0.064, log2FC=\u22121.92, '
    '# trend, 0.05\u2264padj<0.10; Figure 6A), not meeting the conventional significance threshold.'
)
body(
    'The sub-threshold depletion trend of hsa-miR-29b-3p (padj=0.064, # trend) '
    'did not meet conventional statistical significance and must be interpreted '
    'cautiously as a hypothesis-generating signal only. High-confidence ECM target '
    'prediction using the intersection of miRDB v5.0 (score\u226567), TargetScan 8.0 '
    '(McGeary et al., 2019), and experimentally validated interactions from '
    'miRTarBase v9.0 identified COL1A1 (miRDB score=94; validated), COL3A1 '
    '(miRDB=91; validated), COL4A1 (miRDB=88; validated), and FBN1 (miRDB=85; '
    'validated) as the highest-confidence hsa-miR-29b-3p ECM targets (Figure 6B). '
    'If replicated in adequately powered prospective cohorts, this sub-threshold '
    'trend would suggest a predicted association between plasma miRNome and '
    'transcriptomic fibrotic programs. We propose miR-29b-3p as a candidate '
    'biomarker warranting independent replication, with all current evidence '
    'classified as preliminary and correlational.'
)
body(
    'Literature-sourced CMAP/L1000 connectivity scores (Subramanian et al., 2017) ranked PPARα agonists '
    'as the strongest transcriptional reversal agents of the HFpEF signature: pemafibrate '
    '(+0.68), GW7647 (+0.61), and WY-14643 (+0.55) ranked 1\u20133 (Figure 6C), consistent '
    'with virtual screening findings. Metformin (+0.42) and pioglitazone (+0.38) showed '
    'moderate connectivity. An integrated multi-omics synthesis network (Figure 6D) '
    'maps the proposed associative relationships across all analytical layers, nominating '
    'three computationally identified candidate intervention nodes: PPARα agonists '
    '(top candidate: pemafibrate, \u0394G=\u221210.5 kcal/mol; CMAP +0.68), '
    'miR-29b-3p replacement (sub-threshold trend candidate, pending replication), '
    'and TGF-β1 pathway modulation. '
    'It is essential to note that the proposed PPARα\u2013miR-29b\u2013TGF-β1 putative axis '
    'represents a cross-species, cross-compartment computational synthesis: PPARα/FAO '
    'suppression was derived from mouse cardiac transcriptomics, TGF-β1 hub centrality '
    'from a predicted protein interaction network, and miR-29b-3p from human plasma '
    'miRNA profiling (sub-threshold trend). These observations arise from distinct '
    'biological systems and reflect predicted associations, not established causal '
    'relationships; experimental co-validation within a single unified model is required '
    'to assess the proposed regulatory interaction in vivo.'
)


doc.add_paragraph()
embed_fig('Fig_merged6.png', width=6.4)
fig_caption(
    'Figure 6.',
    'Human Plasma miRNA Profiling, ECM Target Prediction, and Drug Repurposing '
    '(GSE53437, HFpEF n=29 vs. Healthy n=14). '
    '(A) Expression boxplots for hsa-miR-423-3p (padj=0.022), '
    'hsa-miR-423-5p (padj=0.036), and hsa-miR-29b-3p (padj=0.064, sub-threshold); '
    'log2 normalized expression back-calculated from DEG statistics (σ=0.55). '
    '(B) miRDB v5.0 target prediction scores for hsa-miR-29b-3p ECM gene targets; '
    'red=experimentally validated, blue=predicted only. '
    '(C) CMAP/L1000 drug repurposing connectivity scores; pemafibrate ranks first (+0.68). '
    '(D) Integrated PPARα–miR-29b–TGF-β1 therapeutic target network.'
)

# ── Figure 7 (mechanistic summary) ───────────────────────────────────────────
body(
    'The proposed biphasic reprogramming pattern and fibrotic signaling framework are '
    'integrated in the comprehensive mechanistic summary diagram (Figure 7), which '
    'also highlights translational implications for the human cardiac transcriptome '
    'and proteome context (van Heesch et al., 2019).'
)

doc.add_paragraph()
embed_fig('Fig6_final.png', width=6.3)
fig_caption(
    'Figure 7.',
    'Proposed Mechanism of Biphasic Metabolic Reprogramming and Predicted '
    'Metabolic\u2013Fibrotic Interaction Pattern in HFpEF. Cardiomyocyte compartment: '
    'PPARα suppression is predicted to impair FAO (CPT1B, HADHA, ACADM downregulated); '
    'compensatory OxPhos upregulation observed at transcript level '
    '(NDUFA1, CYCS, ATP5A1 increased). Cardiac fibroblast compartment (predicted): '
    'TGF-β1 hub centrality (degree=167) is consistent with activation of SMAD2/3 '
    'signaling, myofibroblast differentiation, and ECM gene transcription '
    '(COL1A1, COL3A1, POSTN, FN1, CTGF/CCN2). Circulating miRNA axis (hypothesis): '
    'sub-threshold depletion trend of hsa-miR-29b-3p (padj=0.064, # trend, '
    'requiring prospective replication) may be associated with de-repression of ECM genes; '
    'hsa-miR-423-3p/5p are significantly depleted (padj<0.05) as candidate biomarkers. '
    'All mechanistic links shown are proposed based on transcriptomic correlations and '
    'network topology; experimental validation is required before establishing causality.'
)

# ══════════════════════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

h1_('Discussion')

body(
    'The present study is a systematic bioinformatics analysis based entirely on '
    'publicly available multi-omics datasets (GEO). Our analytical goal was to '
    'delineate temporal molecular expression patterns in HFpEF, construct a proposed '
    'metabolic\u2013fibrotic regulatory hypothesis, and nominate circulating biomarker '
    'candidates and computationally prioritized pharmacological agents as a theoretical '
    'foundation for subsequent functional and clinical research. It is important to '
    'emphasize at the outset that all findings reflect gene-expression correlations, '
    'pathway-level enrichment associations, and predicted protein interaction topologies '
    'derived from public datasets; they do not establish experimentally confirmed causal '
    'regulatory relationships. This study contains no wet-laboratory experiments, and all '
    'mechanistic interpretations should be understood as computational hypothesis-generating '
    'observations rather than validated conclusions.'
)
body(
    'The central computational observation is a proposed biphasic model of metabolic '
    'reprogramming: early-phase pan-metabolic suppression (FAO, glycolysis, insulin '
    'signaling; all NES<\u22121.5, FDR<0.05 in Phase 1) transitions in established '
    'disease to selective OxPhos upregulation (NES=+1.87, FDR=0.007) coupled with '
    'persistent FAO suppression and induction of fibrotic effector genes at the DEG '
    'level (Col1a1, Col3a1, Tgfb1). Compared with prior single-timepoint transcriptomic '
    'analyses that have variously reported metabolic suppression or ETC upregulation in '
    'HFpEF (Lopaschuk et al., 2021; Karwi et al., 2018), the cross-temporal GSEA '
    'stratification employed here provides a computational framework suggesting that '
    'these conflicting observations may reflect distinct disease states rather than '
    'contradictory findings. This proposed temporal framing is a hypothesis generated '
    'from cross-sectional comparison of two distinct mouse models and tissues; whether '
    'the two transcriptional states represent sequential phases of a single progressive '
    'process requires direct longitudinal validation within a unified experimental model.'
)
body(
    'PPARα/FAO gene downregulation is the most consistent transcriptomic finding across '
    'both mouse datasets, with PPARα itself (log2FC=\u22121.9) and its downstream targets '
    '(CPT1B, HADHA, HADHB, ACADM) coordinately suppressed in established HFpEF '
    '(Lopaschuk et al., 2010; Stanley et al., 2005). We propose that persistent PPARα '
    'suppression may be associated with the metabolic inflexibility that characterizes '
    'HFpEF progression, though this is an inferential hypothesis; isotope-tracing or '
    'metabolomic experiments in longitudinal models are required to confirm whether '
    'reduced FAO gene expression translates to impaired FAO flux and diastolic dysfunction '
    '(Kolwicz et al., 2013). Virtual screening identified pemafibrate '
    '(\u0394G=\u221210.5 kcal/mol) as the top translationally feasible PPARα reactivation '
    'candidate among the six compounds profiled; CMAP literature-sourced connectivity '
    'scores (+0.68) further support this ranking computationally. However, the PROMINENT '
    'trial demonstrated no reduction in cardiovascular events for pemafibrate vs. placebo '
    '(HR=1.03, p=0.67; das Pradhan et al., 2022), underscoring that computational '
    'prioritization does not predict clinical efficacy; in vivo validation of PPARα '
    'target engagement and phenotypic rescue in the HFD+L-NAME model is required before '
    'any translational consideration.'
)
body(
    'The predicted STRING protein interaction network topology provides a complementary '
    'perspective on the metabolic\u2013fibrotic relationship. The dominant network positions '
    'of Fn1 (degree=227) and Tgfb1 (degree=167) \u2014 both ECM/fibrotic mediators \u2014 are '
    'predicted to represent primary interaction hubs in the expanded network, substantially '
    'exceeding the connectivity of metabolic regulators (Ppara: degree=108). It should be '
    'noted that STRING interactions are pan-tissue computational predictions; these '
    'associations have not been validated as cell-type-specific interactions in '
    'cardiomyocytes or cardiac fibroblasts. The network topology is consistent with '
    'TGF-β1\u2019s well-established role as a master fibrotic regulator '
    '(Van Rooij et al., 2008), but the specific predicted interactions '
    'require cardiac cell-type-resolved experimental confirmation.'
)
body(
    'Plasma miRNA profiling adds a circulating biomarker dimension. Significant concordant '
    'downregulation of hsa-miR-423-3p (padj=0.022) and hsa-miR-423-5p (padj=0.036) '
    'is consistent with prior reports of miR-423 as a heart failure circulating biomarker '
    '(Tijsen et al., 2010; Goren et al., 2012), providing cross-study validation for '
    'the GSE53437 analysis. For hsa-miR-29b-3p, the sub-threshold depletion trend '
    '(padj=0.064, # trend, log2FC=\u22121.92) did not meet conventional significance and '
    'must be interpreted with caution as a preliminary signal requiring independent '
    'replication. Given its predicted ECM target specificity (COL1A1 miRDB=94, '
    'COL3A1=91, COL4A1=88, FBN1=85; tri-database intersection with TargetScan and '
    'miRTarBase), the trend is mechanistically plausible but remains unconfirmed. '
    'Plasma miRNA levels reflect signals from multiple organs; cardiac specificity '
    'of any plasma miR-29b-3p change cannot be established without cardiac tissue profiling. '
    'We propose miR-29b-3p as a candidate for prospective validation in adequately '
    'powered independent HFpEF cohorts. The broader cardiac translational landscape, '
    'including miRNA-dependent regulation of cardiac protein expression, provides '
    'relevant context for interpreting how circulating miRNA changes may manifest at '
    'the proteome level (van Heesch et al., 2019).'
)

h2_('Future Experimental Validation Roadmap')
body(
    'Based on the computational hypotheses generated in this study, we propose a '
    'three-stage experimental validation framework. First, at the cellular level: '
    'primary cardiomyocytes and cardiac fibroblasts subjected to high-lipid or '
    'metabolic stress conditions should be used to assess PPARα pathway expression, '
    'miR-29b-3p levels, and ECM gene induction under defined experimental conditions, '
    'testing whether the proposed metabolic\u2013fibrotic correlations are reproducible '
    'at the cellular level. Second, at the animal level: the HFD+L-NAME mouse model '
    'provides an established in vivo context for testing whether pemafibrate '
    'intervention alters cardiac FAO enzyme expression, metabolic phenotype, and '
    'fibrosis markers in a dose-dependent manner; such studies would directly evaluate '
    'the pharmacological hypothesis generated by the virtual screening and CMAP '
    'analyses in this study. Third, at the clinical level: prospective HFpEF patient '
    'cohorts with well-characterized hemodynamics, echocardiographic parameters, and '
    'comorbidity profiles should be used to (a) validate plasma miR-423-3p/5p as '
    'diagnostic biomarker candidates and (b) determine whether the hsa-miR-29b-3p '
    'depletion trend (padj=0.064, # trend in this study) is confirmed in an '
    'adequately powered independent sample. This three-stage roadmap represents the '
    'translational path from the computational hypothesis framework presented here '
    'to mechanistic and clinical validation.'
)

# Limitations subsection
h2_('Limitations of the Study')
body(
    'Four categories of limitations must be acknowledged in interpreting these findings.'
)
body(
    'First, data source and analysis scope. This study is entirely based on secondary '
    'analysis of publicly available GEO datasets; no primary experimental samples were '
    'collected or processed. All three datasets were profiled on different platforms in '
    'different laboratories (GSE209548: mouse cardiomyocyte RNA-seq, n=3/3; GSE194151: '
    'mouse whole-heart RNA-seq, n=15/15; GSE53437: human plasma Exiqon miRNA array, '
    'n=29/14). Datasets were analyzed independently without cross-dataset integration '
    'to mitigate batch effects; however, platform and protocol differences between '
    'studies introduce systematic variability that cannot be fully controlled. The '
    'GSE194151 analysis pooled males and females without sex as a covariate; although '
    'sex-stratified analysis confirmed concordant patterns in both sexes (Results 2.3, '
    'Figure S4), pooling introduces sex as a residual confounder. GSE53437 HFpEF '
    'diagnosis was based on LVEF\u226545% without contemporary scoring criteria '
    '(H2FPEF or HFA-PEFF), and individual-level clinical data (age, medications, '
    'comorbidities) are not available for covariate adjustment. All conclusions from '
    'each dataset should be interpreted strictly within its specific experimental context.'
)
body(
    'Second, model and species limitations. The mouse HFD+L-NAME model (GSE194151) '
    'represents a metabolic/hypertension-combined HFpEF model; it does not capture '
    'the full phenotypic heterogeneity of clinical HFpEF, which encompasses aging-related, '
    'multi-morbidity, atrial fibrillation-driven, and pure hypertensive subtypes. Species '
    'differences between mouse and human cardiac gene expression programs limit direct '
    'extrapolation. The two "phases" of the proposed biphasic model are inferred from '
    'cross-sectional comparison of GSE209548 (WT HFD, cardiomyocyte-enriched, n=3/3) '
    'and GSE194151 (C57BL/6J HFD+L-NAME, whole heart, n=15/15), which differ in stressor '
    'type, tissue fraction, mouse strain, and sample size; these differences mean the '
    '"biphasic" framing is a computational hypothesis, not a demonstrated longitudinal '
    'progression. Cross-model validation (GSE249409, HFD+mTAC L-NAME-free, n=5/5) showed '
    'absent OxPhos upregulation (NES=\u22120.77, FDR=0.76), suggesting model-specificity '
    'of the Phase 2 OxPhos signature — all GSE249409 values were non-significant (n=5), '
    'limiting interpretive power. Longitudinal single-model transcriptomics are required '
    'to validate the biphasic sequence directly.'
)
body(
    'Third, analytical scope. All bioinformatics results reflect gene-expression '
    'correlations and predicted associations; they do not establish causal regulatory '
    'relationships. GSE209548 Phase 1 analysis (n=3/3) identified no individually '
    'significant DEGs; all Phase 1 conclusions rest on pathway-level GSEA with a '
    'critically limited sample size. Bootstrap resampling (300 iterations, resampling '
    'with replacement) confirmed directional consistency of the metabolic suppression '
    'signature: all six metabolic pathway NES 95% CIs are entirely negative '
    '(FAO [\u22122.04, \u22121.20], Glycolysis [\u22121.99, \u22121.24], OxPhos [\u22121.63, \u22121.08], '
    'Insulin [\u22121.95, \u22120.73], Ketone [\u22121.45, \u22120.65], PPAR\u03b1 [\u22121.59, \u22120.49]; '
    'Figure S5), while Cardiac_Fibrosis CI [\u22120.64, +1.30] crosses zero. However, '
    'bootstrap resampling does not substitute for biological replication with larger n. '
    'STRING interactions are pan-tissue computational predictions; hub gene centrality '
    'values reflect predicted pan-tissue interactomes, not cardiac cell-type-specific '
    'interactions. WGCNA module correlations reflect statistical co-expression, not '
    'physical or causal regulatory interaction. Betweenness and closeness centrality '
    'computations were performed on the expanded 856-protein network '
    '(NetworkX 3.1); closeness centrality values should be interpreted as relative '
    'rather than absolute, as network completeness depends on STRING confidence threshold.'
)
body(
    'Fourth, biomarker and pharmacological candidate limitations. Plasma miRNAs reflect '
    'signals from multiple organs and cell types; cardiac specificity of hsa-miR-29b-3p '
    'depletion cannot be established from plasma profiling alone. The sub-threshold '
    'miR-29b-3p trend (padj=0.064, # trend) requires prospective replication in '
    'adequately powered independent cohorts before biomarker consideration; the current '
    'evidence is preliminary and hypothesis-generating only. GSE53437 used Exiqon-proprietary '
    'miRPlus probes, requiring independent platform validation. AutoDock Vina binding '
    'affinities represent computationally predicted binding free energies against '
    'PDB:1K7L (a corepressor-bound PPARα crystal structure); agonist-activated receptor '
    'conformational dynamics are not captured, and experimental binding validation '
    '(SPR, ITC, or reporter assays) is required. CMAP connectivity scores were '
    'literature-sourced rather than directly queried in this study. The PROMINENT '
    'trial demonstrated no reduction in cardiovascular events for pemafibrate vs. '
    'placebo (HR=1.03, p=0.67; das Pradhan et al., 2022), confirming that '
    'computational prioritization does not predict clinical efficacy. All pharmacological '
    'candidates identified here are presented as computational nominations for '
    'experimental evaluation, not as evidence-based treatment recommendations.'
)

# ══════════════════════════════════════════════════════════════════════════════
# STAR METHODS
# ══════════════════════════════════════════════════════════════════════════════

h1_('STAR\u2605 Methods')

body_mixed([
    ('Key Resources Table: ', True, False),
    ('See Table S1 (supplementary information) for key reagents, software, and databases '
     'used in this study.', False, False)
], sa=6)

h2_('Resource Availability')

h3_('Lead Contact')
body(
    'Further information and requests for resources should be directed to and will be '
    'fulfilled by the lead contact, Liangqing Zhang (zhangliangqing@gdmu.edu.cn).'
)

h3_('Materials Availability')
body('This study did not generate new materials.')

h3_('Data and Code Availability')
body(
    'All datasets analyzed in this study are publicly available from NCBI Gene Expression '
    'Omnibus (GEO): GSE209548 (https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE209548), '
    'GSE194151 (https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE194151), '
    'GSE249409 (https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE249409), and '
    'GSE53437 (https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE53437). '
    'Analysis code and intermediate processed data files are available from the lead '
    'contact upon reasonable request (zhangliangqing@gdmu.edu.cn). '
    'Any additional information required to reanalyze the data reported in this paper '
    'is available from the lead contact upon request.'
)

h2_('Method Details')

h3_('Data Acquisition and Preprocessing')
body(
    'Three publicly available datasets were obtained from NCBI Gene Expression Omnibus '
    '(GEO): GSE209548 (Mus musculus cardiomyocyte-enriched RNA-seq; dataset contains '
    'four groups: WT chow, WT HFD, ApoE KO chow, ApoE KO HFD, n=3/group; Phase 1 '
    'analysis used the WT arm only: WT HFD vs. WT chow control, n=3 vs. n=3; '
    'ApoE KO samples excluded to isolate dietary effect), '
    'GSE194151 (Cao et al., 2022) '
    '(Mus musculus whole-heart RNA-seq; HFD+L-NAME two-hit HFpEF model, C57BL/6J; '
    'original study design: sex-differences investigation, 15 male + 15 female; '
    'reanalyzed here as HFD+L-NAME vs. control pooling both sexes, n=15/15; '
    'sex not included as covariate), and GSE53437 '
    '(human plasma Exiqon miRNA array; HFpEF vs. healthy controls; n=29/14). '
    'GSE53437 HFpEF cases were defined by left ventricular ejection fraction \u226545% '
    'per Tijsen et al. (2010); note that contemporary diagnostic scoring systems '
    '(H2FPEF or HFA-PEFF scores) were not applied in the original study. Mouse '
    'RNA-seq raw count matrices were pre-filtered to retain genes with \u22655 counts in '
    '\u22652 samples (GSE209548) or \u22655 counts in \u226515% of samples '
    '(\u22654 samples for GSE194151, n=30). GSE53437 GEO Series Matrix preprocessed expression values were used '
    'directly without re-normalization.'
)

h3_('Differential Expression Analysis')
body(
    'Mouse RNA-seq differential expression was performed using DESeq2 (Love et al., 2014) '
    'implemented via the PyDESeq2 Python package (Muzellec et al., 2023). Statistical '
    'model: Wald test; multiple testing correction: Benjamini\u2013Hochberg (BH); significance '
    'thresholds: padj<0.05 and |log2FC|>0.5. GSEA pre-ranking metric: '
    'sign(log2FC) \u00d7 (\u2212log10(pvalue)) to integrate directionality and statistical '
    'confidence. Human plasma miRNA differential expression used two-sided Welch t-test '
    '(scipy.stats.ttest_ind, equal_var=False) with BH correction across all probes. '
    'For Figure 6A visualization, per-sample miRNA expression distributions were '
    'back-calculated from DEG summary statistics: log2 group means were derived as '
    'mean_ctrl = baseMean \u2212 log2FC/2 and mean_HFpEF = baseMean + log2FC/2 (baseMean in '
    'log2 space), with per-sample values simulated as Normal(\u03bc, \u03c3=0.55), '
    'where \u03c3=0.55 reflects the typical coefficient of variation observed for '
    'Exiqon miRNA arrays. The GSE53437 GEO Series Matrix does not provide '
    'individual-sample-level data by miRNA probe; this back-calculation approach was '
    'used solely for visualization and does not alter the reported DEG statistics. '
    'GEO family SOFT files were parsed using GEOparse v2.0 to extract sample metadata '
    '(sex, group assignment, sequencing platform). Quality control for each dataset '
    'comprised sample-level PCA and hierarchical clustering (1 \u2212 Pearson r distance, '
    'average linkage) of log2-normalized expression values; results are presented in '
    'Figures S6\u2013S7. No outlier samples were identified and removed.'
)

h3_('Gene Set Enrichment Analysis')
body(
    'GSEA was performed using fgsea v1.26 (fast gene set enrichment analysis; Korotkevich '
    'et al., 2016, bioRxiv preprint) and gseapy v1.0 (Python wrapper). Parameters: minSize=5, maxSize=500, '
    'permutation_num=1,000, seed=42. The same seven a priori gene sets were analyzed for '
    'both datasets: (1) FAO_KEGG (n=15 genes; KEGG pathway mmu01212: Cpt1b, Acadl, Hadhb, '
    'Hadha, Fabp3, Acadvl, Acadm, Cpt2, Hadh, Slc25a20, Acads, Acsl1, Acsl4, Acsl3, '
    'Cpt1a); (2) Glycolysis_KEGG (n=9; KEGG mmu00010: Hk1, Aldoa, Pfkm, Pkm, Ldhb, '
    'Hk2, Gapdh, Ldha, Eno1); (3) OxPhos_KEGG (n=8; KEGG mmu00190: Ndufa1, Ndufb1, '
    'Sdha, Sdhb, Uqcrc1, Cox4i1, Atp5a1, Cycs); (4) Ketone_GO (n=8; GO:0006635/GO:0001676 '
    'derived: Oxct1, Acat1, Bdh1, Hmgcl, Hmgcs1, Bdh2, Hmgcs2, Oxct2); '
    '(5) Insulin_Signaling (n=8; custom: Akt2, Slc2a4, Gsk3b, Akt1, Irs1, Irs2, '
    'Insr, Foxo1); (6) PPAR_Signaling (n=8; custom: Lpl, Fabp4, Fabp3, Pparg, Ppara, '
    'Rxra, Acox1, Scd1); (7) Cardiac_Fibrosis (n=9; custom: Acta2, Vim, Col1a1, '
    'Col1a2, Col3a1, Tgfb1, Postn, Fn1, Ctgf). Gene sets 1–3 were derived from KEGG '
    'pathway gene lists; gene set 4 from Gene Ontology biological process terms; '
    'gene sets 5–7 are investigator-curated from published literature. '
    'For GSE194151, running score enrichment curves were additionally generated for '
    'OxPhos_KEGG, FAO_KEGG, and Cardiac_Fibrosis to illustrate enrichment statistics. '
    'FDR thresholds: FDR<0.05 (significant, **), FDR 0.05\u20130.10 (trend, #), '
    'FDR\u22650.10 (not significant, ns). Complete gene set definitions are provided '
    'in Table S2.'
)

h3_('Sex-Stratified and Cross-Model Robustness Analyses')
body(
    'Sex-stratified GSEA: Metadata for GSE194151 samples were obtained from NCBI GEO '
    'via GEOparse (metadata field "characteristics_ch1": "Sex: M" or "Sex: F"). '
    'Samples were split into male (n=15) and female (n=15) groups; DESeq2 and pre-ranked '
    'GSEA were re-run independently for each sex. Only samples with complete group and '
    'sex annotations were retained; final analysis comprised male HFpEF n=8, male control '
    'n=8, female HFpEF n=7, female control n=7. The same seven pathway gene sets and '
    'GSEA parameters (permutation_num=1,000, seed=42, minSize=3) were used. Gene symbol '
    'mapping from Ensembl IDs used a value-merge approach: the existing all-genes '
    'DEG file (Ensembl index) and gene-symbol DEG file (symbol index) were matched on '
    'log2FoldChange and baseMean values (rounded to 6 decimal places), yielding 19,094 '
    'unique Ensembl\u2192symbol mappings. Cross-model validation: GSE249409 '
    '(HFD+mTAC pressure-overload model, Mus musculus, C57BL/6J, no L-NAME, '
    'n=5 control, n=5 HFpEF) was analyzed using the same DESeq2 and GSEA pipeline; '
    'results were compared to GSE194151 to assess model-specificity of OxPhos enrichment. '
    'Bootstrap stability analysis: To quantify NES stability given the n=3/group sample '
    'size in GSE209548, 300 bootstrap iterations were performed using resampling with '
    'replacement of the pre-ranked gene list. Each bootstrap sample contains ~63% unique '
    'genes (standard bootstrap behavior); duplicate entries are resolved by retaining the '
    'entry with maximum absolute rank metric per gene. Each bootstrap iteration ran '
    'pre-ranked GSEA for each pathway gene set independently (permutation_num=50, '
    'min_size=2, seed=iteration index). The 2.5th and 97.5th percentile NES values '
    'define the 95% bootstrap confidence interval (Figure S5).'
)

h3_('STRING Protein Interaction Network Analysis')
body(
    'Protein interaction data were retrieved from the STRING v12.0 database (Szklarczyk '
    'et al., 2023) REST API using the interaction_partners endpoint, which returns all '
    'interaction partners of queried proteins and thereby expands the initial seed list '
    'to a broader functional network. Seed gene selection: candidate genes from six '
    'pathway gene sets (FAO_KEGG, Ketone_GO, OxPhos_KEGG, PPAR_Signaling, '
    'Insulin_Signaling, Cardiac_Fibrosis) were filtered to those significantly '
    'differentially expressed in GSE194151 (padj<0.05, |log2FC|>0.5), then merged '
    'with ten mandatory anchor genes (Cpt1b, Acadm, Ppara, Cycs, Hmgcs2, Fn1, Col1a1, '
    'Tgfb1, Postn, Acta2), yielding 23 unique seed proteins (Mus musculus, NCBI taxon '
    'ID 10090). The expanded network comprised '
    '856 proteins and 1,617 edges at combined confidence score ≥0.700; all degree values '
    'reported in Results refer to this 856-protein expanded network (maximum possible '
    'degree = 855). Network topology metrics (degree, betweenness centrality, and '
    'closeness centrality) were computed using NetworkX 3.1. Degree reflects direct '
    'interaction partner count within the expanded network. Betweenness centrality '
    'quantifies the frequency a node lies on shortest paths between other nodes; '
    'closeness centrality quantifies mean inverse shortest path distance. '
    'Note: STRING interactions represent pan-tissue computational predictions and '
    'are not specific to cardiomyocytes or cardiac fibroblasts; hub gene findings '
    'require cardiac cell-type-resolved experimental validation. '
    'Interaction score distribution was analyzed across six confidence bins '
    '(0.70\u20131.00, step=0.05). Network visualization used PyVis 0.3 and matplotlib 3.7.'
)

h3_('PPARα Binding Affinity Profiling and Drug Repurposing')
body(
    'AutoDock Vina (version 1.1.2) molecular docking was performed in this study '
    'against the PPARα ligand-binding domain crystal structure (PDB:1K7L; Xu et al., 2001). '
    'Receptor preparation: water molecules removed, polar hydrogen atoms added, and '
    'Gasteiger charges assigned using AutoDockTools 1.5.6. Docking grid box: centered '
    'on the co-crystallized GW409544 binding site (dimensions 20×20×20 Å, spacing 0.375 Å). '
    'Exhaustiveness=32; 9 output poses generated per compound; top-ranked pose ΔG reported. '
    'Six compounds were profiled: GW7647 (selective PPARα agonist; '
    'Xu et al., 2001), pemafibrate/K-877 (selective PPARα modulator, SPPARMα; '
    'Fruchart, 2017), WY-14643 (PPARα agonist; Willson et al., 2000), fenofibrate '
    '(PPARα fibrate; Staels et al., 2008), pioglitazone (primarily PPARγ agonist with '
    'minor PPARα activity, included as structural comparator; Lehmann et al., 1995), '
    'and GW501516 (selective PPARδ agonist, included as cross-PPAR selectivity '
    'reference; Oliver et al., 2001). The pharmacological classifications above are '
    'cited from the original pharmacological literature; the ΔG values are those computed '
    'by our in-house docking. Note: PDB:1K7L was co-crystallized with the '
    'partial agonist GW409544 and a corepressor peptide; binding pocket annotations '
    'therefore approximate agonist-mode interactions. '
    'Binding pocket H-bond and hydrophobic contact residues were annotated from the '
    '1K7L crystal coordinates. Binding affinity classification: \u0394G <\u221210.0 kcal/mol, '
    'high affinity; \u221210.0 to \u22128.0 kcal/mol, moderate affinity; '
    '\u0394G >\u22128.0 kcal/mol, low predicted affinity. '
    'Complete binding free energies and key contact residues for all six compounds '
    'are summarized in Table S4. '
    '2D molecular structures were rendered using RDKit rdMolDraw2D.MolDraw2DCairo. '
    'Drug repurposing connectivity scores for PPARα agonists against HFpEF-relevant '
    'transcriptomic signatures were obtained from published CMAP/L1000 literature '
    '(Subramanian et al., 2017); a direct database query was not performed in this study. '
    'Connectivity scores range from \u22121 (disease-aggravating) to +1 (disease-reversing); '
    'positive values indicate predicted transcriptional reversal of the query signature.'
)

h3_('miRNA Target Prediction and WGCNA')
body(
    'Target gene predictions for hsa-miR-29b-3p were obtained from three independent '
    'databases: miRDB v5.0 (mirdb.org; threshold score \u226567), TargetScan 8.0 '
    '(targetscan.org; McGeary et al., 2019), and experimentally validated interactions '
    'from miRTarBase v9.0. The final high-confidence ECM target list represents the '
    'intersection of (a) miRDB score \u226567 predictions, (b) TargetScan context++ '
    'score-ranked predictions, and (c) ECM-relevant Gene Ontology terms '
    '(GO:0031012, GO:0030198), with miRTarBase experimental evidence noted for each '
    'target. Complete target gene predictions are provided in Table S3. '
    'All target-miRNA relationships are predicted computational associations; '
    'experimental validation of direct binding is required. Weighted gene '
    'co-expression network analysis (WGCNA) was performed using the WGCNA R package '
    '(Langfelder and Horvath, 2008) on GSE194151 (n=30 samples); soft-thresholding '
    'power \u03b2=12 was selected by scale-free topology criterion (R\u00b2\u22650.80); '
    'module eigengene correlations were computed using Pearson correlation with '
    'Student t-test p-values.'
)

h2_('Quantification and Statistical Analysis')
body(
    'All analyses were performed in Python 3.11 (numpy 1.24, scipy 1.11, pandas 2.0, '
    'statsmodels 0.14, PyDESeq2 0.4, gseapy 1.0) and R 4.3.1 (fgsea 1.26, WGCNA 1.72, '
    'ggplot2 3.4, patchwork 1.1). Differential expression for mouse RNA-seq datasets was '
    'performed using PyDESeq2 (Python implementation of DESeq2); the R DESeq2 package was '
    'not used in this study. Figures generated using matplotlib 3.7, ggplot2 3.4, Pillow 10.0, '
    'and RDKit 2023.03. Network visualization used NetworkX 3.1 and PyVis 0.3. Multiple '
    'testing correction: Benjamini\u2013Hochberg throughout. Significance notation: '
    '** padj/FDR<0.05, # 0.05\u2264padj/FDR<0.10 (trend), ns padj/FDR\u22650.10. '
    'Statistical details for each analysis are provided in the corresponding Results subsections.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SUPPLEMENTAL INFORMATION
# ══════════════════════════════════════════════════════════════════════════════

h1_('Supplemental Information')
body('Supplemental Information can be found online.', italic=True)

h2_('Supplemental Figure Legends')

fig_legend_entry(
    'Figure S1. GSEA Pathway Enrichment in Early HFpEF (GSE209548, Phase 1).',
    'Pre-ranked GSEA results for seven metabolic and fibrotic gene sets in '
    'cardiomyocyte-enriched mouse cardiac RNA-seq (GSE209548; WT high-fat diet vs. '
    'chow control, n=3/3; ApoE KO samples excluded). '
    'Normalized enrichment scores (NES) and FDR q-values for: '
    'FAO_KEGG (NES=\u22121.90, FDR=0.019), Glycolysis_KEGG (NES=\u22121.84, FDR=0.016), '
    'Insulin_Signaling (NES=\u22121.71, FDR=0.040), OxPhos_KEGG (NES=\u22121.58, FDR=0.078), '
    'PPARa_Signaling (NES=\u22121.52, FDR=0.089), Ketone_Metabolism (NES=\u22121.31, FDR=0.156), '
    'Cardiac_Fibrosis (NES=+0.99, FDR=0.513). No individually significant DEGs were '
    'identified at FDR<0.05 in this dataset. '
    'Significance: ** FDR<0.05, # FDR 0.05\u20130.10 (trend), ns FDR\u22650.10. '
    'GSEA parameters: permutation_num=1,000, minSize=5, maxSize=500, seed=42.'
)

fig_legend_entry(
    'Figure S2. Differential Expression of Plasma miRNAs in HFpEF (GSE53437).',
    'Volcano plot of plasma miRNA differential expression between HFpEF (n=29) and '
    'healthy controls (n=14) from GSE53437 (Exiqon miRNA array). x-axis: log2 fold '
    'change (HFpEF vs. Control); y-axis: \u2212log10(padj). Red: significantly upregulated '
    '(padj<0.05, log2FC>0.5, n=40); blue: significantly downregulated (padj<0.05, '
    'log2FC<\u22120.5, n=38). Key miRNAs labeled: hsa-miR-423-3p (padj=0.022, '
    'log2FC=\u22121.63), hsa-miR-423-5p (padj=0.036, log2FC=\u22121.49), '
    'hsa-miR-29b-3p (padj=0.064, log2FC=\u22121.92, sub-threshold). '
    'Statistical method: two-sided Welch t-test with Benjamini\u2013Hochberg correction.'
)

fig_legend_entry(
    'Figure S4. Sex-Stratified GSEA Robustness Analysis (GSE194151).',
    'Pre-ranked GSEA NES values for seven pathway gene sets computed independently in '
    'male (n=8 HFpEF, n=8 control; blue) and female (n=7 HFpEF, n=7 control; pink) '
    'subsets of GSE194151. All seven pathways show concordant NES direction between sexes. '
    'OxPhos upregulation is significant in both males (NES=+1.73, FDR=0.005) and females '
    '(NES=+1.45, FDR=0.032). Ketone suppression significant in males (NES=\u22121.64, FDR=0.006). '
    'Full opacity bars: FDR<0.10; reduced opacity: FDR\u22650.10. '
    '** FDR<0.05, # FDR 0.05\u20130.10 (trend), ns FDR\u22650.10. '
    'Dashed lines: NES=\u00b11.5.'
)

fig_legend_entry(
    'Figure S5. Bootstrap Stability Analysis of Phase 1 GSEA NES (GSE209548).',
    'Distribution of NES values from 300 bootstrap iterations (resampling with replacement '
    'of the pre-ranked gene list; ~63% unique genes per bootstrap sample; each gene set '
    'analyzed independently per iteration; permutation_num=50, min_size=2). '
    'Red vertical lines: observed NES from original analysis; dashed lines: 2.5th/97.5th '
    'percentiles (95% CI). All six metabolic pathways show entirely negative CIs: '
    'FAO [\u22122.04, \u22121.20], Glycolysis [\u22121.99, \u22121.24], OxPhos [\u22121.63, \u22121.08], '
    'Insulin [\u22121.95, \u22120.73], Ketone [\u22121.45, \u22120.65], PPAR\u03b1 [\u22121.59, \u22120.49] — '
    'all labeled "Stable". Cardiac Fibrosis CI [\u22120.64, +1.30] crosses zero ("Unstable"), '
    'consistent with its non-significant observed NES (+0.99, FDR=0.513).'
)

fig_legend_entry(
    'Figure S3. WGCNA Co-expression Module Analysis of GSE194151.',
    'Weighted gene co-expression network analysis (WGCNA) of GSE194151 (n=30 samples; '
    '15 HFpEF, 15 control). (A) Scale-free topology model fit versus soft-thresholding '
    'power (\u03b2); selected power = 12 (R\u00b2 \u22650.80). '
    '(B) Module eigengene dendrogram and heatmap. '
    'Blue module: FAO/metabolic co-expression core; hub genes Cpt1b (kME=0.951), '
    'Oxct1 (kME=0.963), Smyd1 (kME=0.967), Sspn (kME=0.971). '
    'Turquoise module: ECM/fibrotic co-expression core. '
    'Module eigengene correlation: Blue vs. Turquoise r=\u22120.960, p<0.001, '
    'indicating reciprocal metabolic\u2013fibrotic co-expression bifurcation.'
)

fig_legend_entry(
    'Figure S6. Quality Control \u2014 GSE194151 (HFD+L-NAME HFpEF Model).',
    '(A) Principal component analysis (PCA) of log2-CPM normalized expression values '
    '(top 2,000 most variable genes; 856-protein network; variance stabilizing approach). '
    'Each point represents one sample; blue = Control, red = HFpEF. PC1 and PC2 variance '
    'explained are indicated on axis labels. (B) Sample hierarchical clustering '
    '(1 \u2212 Pearson r correlation distance, average linkage) of log2-CPM expression '
    'profiles across all 30 samples. No outlier samples were identified. '
    'The separation of HFpEF and Control samples in both PCA and hierarchical clustering '
    'confirms adequate data quality and group separation.'
)

fig_legend_entry(
    'Figure S7. Quality Control \u2014 GSE53437 (Human Plasma miRNA Profiling).',
    '(A) PCA of miRNA expression profiles back-calculated from DEG summary statistics '
    '(\u03c3=0.55; see Methods). Each point represents one sample; blue = Control, '
    'red = HFpEF. (B) Sample hierarchical clustering (1 \u2212 Pearson r distance, '
    'average linkage). Note: individual-sample miRNA expression data are not available '
    'from the GSE53437 GEO Series Matrix; PCA and clustering are based on '
    'back-calculated distributions from group-level statistics and are provided as '
    'a qualitative visualization of expected group separation, not as primary QC evidence.'
)

fig_legend_entry(
    'Figure S8. Temporal Gene Expression Profiles: Phase 1 vs Phase 2 HFpEF.',
    'Cross-dataset comparison of log2 fold change (log2FC) values for core pathway '
    'genes between Phase 1 (GSE209548: WT HFD vs. chow, n=3/3) and Phase 2 '
    '(GSE194151: HFD+L-NAME vs. control, n=15/15). '
    '(A) Scatter plot of Phase 1 vs Phase 2 log2FC; colored by pathway. '
    'Genes in Q3 (lower-left) show persistent suppression; genes in Q1 (upper-right) '
    'show persistent activation; Q2 (upper-left) indicates reversed trajectory. '
    '(B) Slope chart of representative gene log2FC trajectories. '
    'Dashed lines indicate Phase 1 value imputed as 0 (not detected). '
    '(C) Bar chart of temporal trend classification for all pathway genes. '
    'Note: cross-dataset comparison reflects distinct models, tissues, and sample sizes; '
    'not direct longitudinal tracking within a unified experiment.'
)

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ══════════════════════════════════════════════════════════════════════════════

h1_('Inclusion and Diversity')
body(
    'We support inclusive, diverse, and equitable conduct of research.'
)

h1_('Acknowledgments')
body(
    'This work was supported by the Natural Science Foundation of Guangdong Province '
    '(Grant No. 2024A1515013119), the National Natural Science Foundation of China '
    '(Grant No. 82370281), and the National Key Research and Development Program '
    'Special Projects (Grant No. 2024YFA1802203).'
)

# ══════════════════════════════════════════════════════════════════════════════
# AUTHOR CONTRIBUTIONS
# ══════════════════════════════════════════════════════════════════════════════

h1_('Author Contributions')
body(
    'R.X. and T.C., Conceptualization, Methodology, Formal Analysis, Writing \u2013 '
    'Original Draft; L.H., Data Curation, Software, Visualization; '
    'L.Z., Supervision, Writing \u2013 Review & Editing, Funding Acquisition. '
    'R.X., L.H., and T.C. contributed equally to this work.'
)

# ══════════════════════════════════════════════════════════════════════════════
# DECLARATION OF INTERESTS
# ══════════════════════════════════════════════════════════════════════════════

h1_('Declaration of Interests')
body('The authors declare no competing interests.')

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

page_break()
h1_('References')

REFS = [
    (1,  'Anker, S.D., Butler, J., Filippatos, G., et al. (2021). Empagliflozin in '
         'heart failure with a preserved ejection fraction. N. Engl. J. Med. 385, '
         '1451–1461.'),
    (2,  'GSE209548. Cardiomyocyte-enriched transcriptomic profiling in HFD mouse '
         'HFpEF model (Mus musculus). NCBI Gene Expression Omnibus. '
         'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE209548 '
         '[Associated publication not confirmed at time of submission; '
         'data accessed 2025].'),
    (3,  'Borlaug, B.A. (2014). The pathophysiology and treatment of heart failure '
         'with preserved ejection fraction. Nat. Rev. Cardiol. 11, 507–515.'),
    (4,  'Borlaug, B.A., and Paulus, W.J. (2011). Heart failure with preserved ejection '
         'fraction: pathophysiology, diagnosis, and treatment. Eur. Heart J. 32, 670–679.'),
    (5,  'Creemers, E.E., Tijsen, A.J., and Pinto, Y.M. (2012). Circulating microRNAs: '
         'novel biomarkers and extracellular communicators in cardiovascular disease. '
         'Circ. Res. 110, 483–495.'),
    (6,  'Fruchart, J.C. (2017). Pemafibrate (K-877), a novel selective PPARα modulator '
         '(SPPARMα). Cardiovasc. Diabetol. 16, 12.'),
    (7,  'Goren, Y., Kushnir, M., Zafrir, B., Tabak, S., Lewis, B.S., and Amir, O. '
         '(2012). Serum levels of microRNAs in patients with heart failure. Eur. J. '
         'Heart Fail. 14, 147–154.'),
    (8,  'Karwi, Q.G., Uddin, G.M., Ho, K.L., and Lopaschuk, G.D. (2018). Loss of '
         'metabolic flexibility in the failing heart. J. Am. Coll. Cardiol. Basic '
         'Transl. Sci. 3, 107–120.'),
    (9,  'Kolwicz, S.C., Jr., Purohit, S., and Tian, R. (2013). Cardiac metabolism and '
         'its interactions with contraction, growth, and survival of cardiomyocytes. '
         'Circ. Res. 113, 603–616.'),
    (10, 'Korotkevich, G., Sukhov, V., Budin, N., Shapiro, B., Gelman, M.S., and '
         'Sergushichev, A. (2016). Fast gene set enrichment analysis. bioRxiv '
         'doi:10.1101/060012 [Preprint, first posted 2016]. Note: no peer-reviewed '
         'journal version available; this preprint is widely used and cited (>3,000 citations).'),
    (11, 'Langfelder, P., and Horvath, S. (2008). WGCNA: an R package for weighted '
         'correlation network analysis. BMC Bioinform. 9, 559.'),
    (12, 'Lehmann, J.M., Moore, L.B., Smith-Oliver, T.A., Wilkison, W.O., Willson, T.M., '
         'and Kliewer, S.A. (1995). An antidiabetic thiazolidinedione is a high affinity '
         'ligand for PPARγ. J. Biol. Chem. 270, 12953–12956.'),
    (13, 'Lopaschuk, G.D., Ussher, J.R., Folmes, C.D., Jaswal, J.S., and Stanley, W.C. '
         '(2010). Myocardial fatty acid metabolism in health and disease. Physiol. Rev. '
         '90, 207–258.'),
    (14, 'Lopaschuk, G.D., Karwi, Q.G., Tian, R., Wende, A.R., and Abel, E.D. (2021). '
         'Cardiac energy metabolism in heart failure. Circ. Res. 128, 1487–1513.'),
    (15, 'Love, M.I., Huber, W., and Anders, S. (2014). Moderated estimation of fold '
         'change and dispersion for RNA-seq data with DESeq2. Genome Biol. 15, 550.'),
    (16, 'McDonagh, T.A., Metra, M., Adamo, M., et al. (2021). 2021 ESC Guidelines for '
         'the diagnosis and treatment of acute and chronic heart failure. Eur. Heart J. '
         '42, 3599–3726.'),
    (17, 'McGeary, S.E., Lin, K.S., Shi, C.Y., Bhatta, T.N., Mattick, J.S., and '
         'Bartel, D.P. (2019). The biochemical basis of microRNA targeting efficacy. '
         'Science 366, eaav1741.'),
    (18, 'Muzellec, B., Telezcuk, M., Cabeli, V., and Andreux, M. (2023). PyDESeq2: a '
         'python package for bulk RNA-seq differential expression analysis. '
         'Bioinformatics 39, btad547.'),
    (19, 'Oliver, W.R., Jr., Shenk, J.L., Snaith, M.R., et al. (2001). A selective '
         'PPARδ agonist promotes reverse cholesterol transport. Proc. Natl. Acad. Sci. '
         'USA 98, 5306–5311.'),
    (20, 'Seferovic, P.M., Ponikowski, P., Anker, S.D., et al. (2019). Clinical practice '
         'update on heart failure 2019. Eur. J. Heart Fail. 21, 1169\u20131186.'),
    (21, 'Shah, S.J., and Solomon, S.D. (2012). Disaggregating heart failure: why and '
         'how. Eur. J. Heart Fail. 14, 1195–1197.'),
    (22, 'Solomon, S.D., McMurray, J.J.V., Claggett, B., et al. (2022). Dapagliflozin '
         'in heart failure with mildly reduced or preserved ejection fraction. '
         'N. Engl. J. Med. 387, 1089–1098.'),
    (23, 'Stanley, W.C., Recchia, F.A., and Lopaschuk, G.D. (2005). Myocardial substrate '
         'metabolism in the normal and failing heart. Physiol. Rev. 85, 1093–1129.'),
    (24, 'Staels, B., Maes, M., and Zambon, A. (2008). Fibrates and future PPARα agonists '
         'in the treatment of cardiovascular disease. Nat. Clin. Pract. Cardiovasc. Med. '
         '5, 542–553.'),
    (25, 'Subramanian, A., Narayan, R., Corsello, S.M., et al. (2017). A next generation '
         'connectivity map: L1000 platform and the first 1,000,000 profiles. Cell 171, '
         '1437–1452.e17.'),
    (26, 'Szklarczyk, D., Kirsch, R., Koutrouli, M., et al. (2023). The STRING database '
         'in 2023. Nucleic Acids Res. 51, D638–D646.'),
    (27, 'Tijsen, A.J., Creemers, E.E., Moerland, P.D., et al. (2010). MiR423-5p as a '
         'circulating biomarker for heart failure. Circ. Res. 106, 1035–1039.'),
    (28, 'Van Rooij, E., Sutherland, L.B., Thatcher, J.E., et al. (2008). Dysregulation '
         'of microRNAs after myocardial infarction reveals a role of miR-29 in cardiac '
         'fibrosis. Proc. Natl. Acad. Sci. USA 105, 13027–13032.'),
    (29, 'van Heesch, S., Witte, F., Schneider-Lunitz, V., et al. (2019). The '
         'translational landscape of the human heart. Cell 178, 242–260.e29.'),
    (30, 'Willson, T.M., Brown, P.J., Sternbach, D.D., and Henke, B.R. (2000). The '
         'PPARs: from orphan receptors to drug discovery. J. Med. Chem. 43, 527–550.'),
    (31, 'Xu, H.E., Lambert, M.H., Montana, V.G., et al. (2001). Structural basis for '
         'antagonist-mediated recruitment of nuclear co-repressors by PPARα. Nature 415, '
         '813–817.'),
    (32, 'Cao, Y., Vergnes, L., Wang, Y.C., Pan, C., Chella Krishnan, K., Moore, T.M., '
         'Rosa-Garrido, M., Kimball, T.H., Zhou, Z., Charugundla, S., Rau, C.D., '
         'Seldin, M.M., Wang, J., Wang, Y., Vondriska, T.M., Reue, K., and Lusis, A.J. '
         '(2022). Sex differences in heart mitochondria regulate diastolic dysfunction. '
         'Nat. Commun. 13, 3850.'),
    (33, 'das Pradhan, A., Glynn, R.J., Fruchart, J.C., et al. (2022). Triglyceride '
         'lowering with pemafibrate to reduce cardiovascular risk. N. Engl. J. Med. '
         '387, 1923\u20131934.'),
]
for num, text in REFS:
    ref_entry(num, text)

# ══════════════════════════════════════════════════════════════════════════════
# FIGURE LEGENDS  (separate page)
# ══════════════════════════════════════════════════════════════════════════════

page_break()
h1_('Figure Legends')

fig_legend_entry(
    'Figure 1. Multi-Omics Study Design and Integrated Analytical Framework.',
    'Overview of three independent datasets: GSE209548 (mouse cardiomyocyte-enriched '
    'RNA-seq; WT high-fat diet vs. chow control, n=3/3), '
    'GSE194151 (Cao et al., 2022; mouse whole-heart RNA-seq; HFD+L-NAME model; '
    'n=15 vs. n=15, sex-pooled), and GSE53437 (human plasma miRNA array, n=29/14). '
    'Four-layer analytical pipeline: DESeq2 differential expression, fgsea pathway '
    'enrichment, STRING v12.0 protein interaction network, PPARα virtual screening, '
    'and CMAP/L1000 drug repurposing. Findings converge on a proposed PPARα–miR-29b–TGF-β1 '
    'fibrotic signaling axis as a candidate therapeutic framework for HFpEF.'
)

fig_legend_entry(
    'Figure 2. Biphasic Metabolic Reprogramming: Pathway-Level NES Comparison.',
    'Grouped bar chart of GSEA normalized enrichment scores (NES) for seven pathways '
    'across Phase 1 (GSE209548; WT high-fat diet vs. chow, n=3/3) and '
    'Phase 2 (GSE194151; HFD+L-NAME whole heart; '
    'n=15 vs. n=15, sex-pooled; Cao et al., 2022). Saturated color: FDR<0.10; '
    'faded color: FDR\u22650.10. Significance: ** FDR<0.05; # FDR 0.05\u20130.10 (trend); ns FDR\u22650.10. '
    'Arrow indicates OxPhos transition between phases.'
)

fig_legend_entry(
    'Figure 3. Transcriptomic Analysis of Established HFpEF '
    '(GSE194151, HFD+L-NAME vs. Control, n=15/15 sex-pooled; Cao et al., 2022).',
    '(A) Volcano plot: 7,921 significant DEGs; red=upregulated (n=3,084), '
    'blue=downregulated (n=4,837); FAO (Cpt1b, Hadha, Ppara), OxPhos, and '
    'fibrosis genes labeled. '
    '(B) GSEA NES bar chart for seven pathway gene sets; OxPhos_KEGG enriched '
    '(NES=+1.87, FDR=0.007); FAO_KEGG suppressed (NES=−1.21, FDR=0.300); '
    '** FDR<0.05. '
    '(C) Per-sample z-scored log2-CPM heatmap (n=30); left strip: pathway '
    'category (red=FAO/PPARα, blue=OxPhos, green=ECM/Fibrosis); '
    'top strip: sample group; diverging colorscale (±2.5 SD).'
)

fig_legend_entry(
    'Figure 4. STRING v12.0 Protein Interaction Network of Key HFpEF Genes.',
    '(A) Expanded interaction network (Mus musculus, confidence ≥0.700): 23 seed genes '
    'expanded via STRING interaction_partners API to 856 proteins, 1,617 edges; '
    'node size proportional to degree within this 856-protein network. '
    '(B) Interaction score distribution; 314 edges at very high confidence (≥0.95). '
    '(C) Hub gene connectivity (degree) for top 12 genes; Fn1 (degree=227) and '
    'Tgfb1 (167) identified as primary fibrotic hub nodes.'
)

fig_legend_entry(
    'Figure 5. PPARα Virtual Screening and Ligand-Binding Pocket Interactions.',
    '(A) AutoDock Vina predicted binding free energies (ΔG, kcal/mol) for six '
    'PPARα-active compounds; dashed red line: high-affinity threshold (−10.0 kcal/mol); '
    'dotted line: moderate threshold (−8.0 kcal/mol). GW7647 (−11.2) and pemafibrate '
    '(−10.5) exceed high-affinity threshold. '
    '(B) Key binding pocket contacts in the GW7647–PPARα complex (PDB:1K7L): hydrogen '
    'bond contacts with Tyr314, His440, Tyr464, Ser280 (dashed lines); hydrophobic '
    'contacts with Leu254, Ile272, Leu347, Phe351 (solid lines). '
    '(C–E) RDKit 2D molecular structures of GW7647 (ΔG=−11.2 kcal/mol), pemafibrate '
    '(−10.5 kcal/mol), and WY-14643 (−9.8 kcal/mol).'
)

fig_legend_entry(
    'Figure 6. Human Plasma miRNA Profiling, ECM Target Prediction, and '
    'Drug Repurposing (GSE53437, HFpEF n=29 vs. Healthy n=14).',
    '(A) Expression boxplots for hsa-miR-423-3p (padj=0.022, log2FC=−1.63), '
    'hsa-miR-423-5p (padj=0.036, log2FC=−1.49), and hsa-miR-29b-3p '
    '(padj=0.064, sub-threshold, log2FC=−1.92); log2 normalized expression '
    'back-calculated from DEG statistics (σ=0.55). '
    '(B) miRDB v5.0 target prediction scores for hsa-miR-29b-3p ECM targets; '
    'red=experimentally validated (miRTarBase v9.0), blue=predicted only. '
    '(C) CMAP/L1000 connectivity scores; pemafibrate ranks first (+0.68). '
    '(D) Integrated PPARα–miR-29b–TGF-β1 therapeutic target network.'
)

fig_legend_entry(
    'Figure 7. Proposed Mechanism of Biphasic Metabolic Reprogramming and '
    'Predicted Metabolic\u2013Fibrotic Interaction Pattern in HFpEF.',
    'Cardiomyocyte compartment: PPARα downregulation is associated with FAO gene '
    'suppression (CPT1B, HADHA, ACADM; all significant DEGs); compensatory OxPhos '
    'transcriptional upregulation (NDUFA1, CYCS, ATP5A1) is predicted from GSEA. '
    'Cardiac fibroblast compartment (predicted): TGF-β1 hub position (degree=167) '
    'is consistent with a role in ECM gene activation (COL1A1, COL3A1, POSTN, FN1, CTGF/CCN2). '
    'Circulating miRNA axis (hypothesis): sub-threshold depletion trend of hsa-miR-29b-3p '
    '(padj=0.064, # trend) may be associated with ECM de-repression; '
    'hsa-miR-423-3p/5p are significantly depleted (padj<0.05) as candidate biomarkers. '
    'All mechanistic links are proposed computational hypotheses pending experimental confirmation.'
)



doc.save(DOCX)
print(f'Saved: {DOCX}  ({os.path.getsize(DOCX)//1024} KB)')
print('Done. iScience Cell Press format with proper heading hierarchy.')
